from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables" / "course_materials"
IMG_DIR = OUT_DIR / "images"


@dataclass
class PublicResource:
    name: str
    url: str
    category: str
    use_case: str
    note: str


PUBLIC_RESOURCES = [
    PublicResource(
        name="OpenAI Platform Docs - Agents",
        url="https://platform.openai.com/docs/guides/agents",
        category="官方学习资料",
        use_case="理解 agent 的基本构成、工具调用与多步骤执行思路",
        note="适合做 agent 基础认知，不适合直接当作变现课程本体。",
    ),
    PublicResource(
        name="IBM Think - Large language models",
        url="https://www.ibm.com/think/topics/large-language-models",
        category="官方学习资料",
        use_case="解释什么是大模型、为什么能做生成与推理",
        note="适合做大众认知解释，语言相对友好。",
    ),
    PublicResource(
        name="blacktwist/social-media-skills",
        url="https://github.com/blacktwist/social-media-skills",
        category="公开 GitHub skill",
        use_case="快速做内容策略、文案写作、复盘分析 demo",
        note="优势是上手快；短板是效果依赖输入质量，不能替代内容训练。",
    ),
    PublicResource(
        name="jihe520/social-push",
        url="https://github.com/jihe520/social-push",
        category="公开 GitHub workflow",
        use_case="演示多平台内容暂存、草稿发布和社媒推送 workflow",
        note="适合展示流程自动化，不适合承诺最终流量结果。",
    ),
    PublicResource(
        name="autoclaw-cc/xiaohongshu-mcp-skills",
        url="https://github.com/autoclaw-cc/xiaohongshu-mcp-skills",
        category="公开 GitHub skill",
        use_case="演示小红书搜索、选题、发布、互动等动作",
        note="适合做平台 demo；真正账号增长仍要靠内容和反馈迭代。",
    ),
    PublicResource(
        name="agent-browser",
        url="https://www.npmjs.com/package/agent-browser",
        category="公开浏览器自动化工具",
        use_case="演示浏览器自动化、页面操作和半自动执行流程",
        note="适合说明 agent 可以操作网页，但不等于完整业务自动化。",
    ),
]


CORE_COURSE_MODULES = [
    {
        "title": "模块 1：先建立正确认知",
        "points": [
            "什么是 AI：用数据、规则和模型帮助人完成识别、生成、判断与辅助决策。",
            "什么是大模型：能够处理自然语言、图片等复杂输入的大规模通用模型。",
            "什么是 agent：能理解目标、调用工具、分步执行、根据结果继续行动的 AI 形态。",
            "什么是 workflow：把一件事拆成稳定步骤，再决定哪些步骤由 AI 协助。",
            "什么是 skill：把某类任务的提示规则、工具调用和输出格式打包成可复用能力。",
        ],
    },
    {
        "title": "模块 2：快速做 demo 的公开工具链",
        "points": [
            "内容类 demo：选题、提纲、标题、改写、整理资料。",
            "发布类 demo：草稿暂存、平台路由、浏览器半自动操作。",
            "研究类 demo：平台搜索、竞品观察、评论整理、需求归纳。",
            "限制要讲透：demo 可以让人看到可能性，但无法代替长期训练、审美、选题判断和复盘。",
        ],
    },
    {
        "title": "模块 3：三类人群的学习路径",
        "points": [
            "新手先过认知关，不要一上来囤工具。",
            "内容创作者先搭一条能反复跑的内容 workflow。",
            "想变现的人先打通“内容 - 私域 - 低客单 - 训练营”的最小闭环。",
        ],
    },
    {
        "title": "模块 4：为什么一定要有训练营",
        "points": [
            "公开资料负责打开认知，训练营负责把动作练熟。",
            "公开 skill 负责快速出样，训练营负责改输入、改流程、改标准。",
            "没有反馈和作业机制，多数人只能看到 demo，看不到稳定结果。",
        ],
    },
]


PERSONA_TRACKS = {
    "newbie": {
        "title": "给新手的 AI 认知手册",
        "subtitle": "适合刚接触 AI、想先搞懂概念与路线的人",
        "positioning": "先学会判断什么值得学，再决定学哪个工具。",
        "pain_points": [
            "看了很多 AI 内容，但概念混在一起。",
            "容易误把工具收藏当作学习进步。",
            "不知道从哪里开始，担心被割韭菜。",
        ],
        "goals": [
            "建立 AI / 大模型 / agent / workflow / skill 的基础地图。",
            "知道什么适合先学，什么不需要现在就学。",
            "形成 14 天的低负担入门路线。",
        ],
        "modules": [
            "先理解概念，不急着装一堆工具。",
            "先学会提问与拆步骤，再学自动化。",
            "先跑 1 条简单 workflow，再考虑更多平台与更多 skill。",
        ],
        "offer": "适合承接到《AI 入门认知训练营》或 69-199 的入门微课。",
    },
    "creator": {
        "title": "给内容创作者的 AI 提效手册",
        "subtitle": "适合做公众号、小红书、朋友圈、短视频文案的人",
        "positioning": "不是让 AI 一次写完，而是把 AI 塞进固定内容流程里。",
        "pain_points": [
            "写内容慢，灵感不稳定。",
            "AI 写出来的东西味儿重、不像自己。",
            "会一点工具，但内容效率没有真正提升。",
        ],
        "goals": [
            "搭起一条从选题到发布的内容 workflow。",
            "知道哪些步骤适合 AI，哪些步骤必须自己判断。",
            "把 AI 从写稿器改成内容协作器。",
        ],
        "modules": [
            "内容 workflow：选题、标题、提纲、初稿、改写、配图、发布前检查。",
            "公开工具链 demo：social-media-skills、social-push、agent-browser。",
            "训练营升级点：反馈、案例拆解、作业批改、账号策略。",
        ],
        "offer": "适合承接到《AI 内容提效陪跑营》或 9.9-49.9 的模板包。",
    },
    "monetizer": {
        "title": "给副业变现人群的 AI 工作流手册",
        "subtitle": "适合想把 AI 变成流量入口、私域产品与服务的人",
        "positioning": "先打通最小成交闭环，再谈大而全课程和高自动化。",
        "pain_points": [
            "看了很多赚钱故事，但不会落到自己的路径上。",
            "不知道是该先卖资料、卖课还是做服务。",
            "担心学了很多工具，结果没有变现闭环。",
        ],
        "goals": [
            "理解“免费资料 - 私域 - 低客单 - 训练营”的产品阶梯。",
            "知道公开 workflow 能做哪些 demo，哪些结果必须靠陪跑。",
            "拿到一个可讲、可卖、可升级的最小方案。",
        ],
        "modules": [
            "变现路径：内容获客、私域承接、产品阶梯、训练营成交。",
            "工具角色：AI 负责提速与打样，人负责价值设计与转化。",
            "升级路径：从资料包到微课，再到系统课、训练营与服务。",
        ],
        "offer": "适合承接到《AI 轻变现训练营》或系统课程咨询。",
    },
}


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    IMG_DIR.mkdir(parents=True, exist_ok=True)


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        test = current + ch
        if draw.textlength(test, font=font) <= max_width or not current:
            current = test
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_card(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    bullets: Iterable[str],
    fill: tuple[int, int, int],
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=24, fill=fill)
    title_font = get_font(34, bold=True)
    body_font = get_font(24)
    draw.text((x1 + 28, y1 + 22), title, font=title_font, fill=(24, 35, 52))
    y = y1 + 82
    for bullet in bullets:
        lines = wrap_text(draw, bullet, body_font, x2 - x1 - 70)
        for idx, line in enumerate(lines):
            prefix = "• " if idx == 0 else "  "
            draw.text((x1 + 32, y), prefix + line, font=body_font, fill=(51, 65, 85))
            y += 34
        y += 10


def build_diagram_concept_map() -> Path:
    path = IMG_DIR / "01_ai_concept_map.png"
    image = Image.new("RGB", (1600, 900), (248, 245, 239))
    draw = ImageDraw.Draw(image)
    title_font = get_font(46, bold=True)
    body_font = get_font(24)
    draw.text((70, 48), "AI / 大模型 / Agent / Workflow / Skill 关系图", font=title_font, fill=(34, 44, 61))
    draw.text((70, 112), "这张图用来帮助新手先建立认知顺序：先懂层级，再懂工具。", font=body_font, fill=(97, 109, 126))

    boxes = [
        ((90, 210, 520, 365), "AI", ["最上层概念", "包括识别、生成、预测、辅助判断"], (255, 228, 196)),
        ((585, 210, 1015, 365), "大模型", ["AI 里最强的一类通用生成能力", "擅长语言、图片、推理与多任务"], (225, 239, 255)),
        ((1080, 210, 1510, 365), "Agent", ["会理解目标", "会调工具", "会按结果继续执行"], (222, 244, 229)),
        ((345, 480, 775, 680), "Workflow", ["把任务拆成稳定步骤", "决定哪步让 AI 协助，哪步由人判断"], (255, 244, 214)),
        ((840, 480, 1270, 680), "Skill", ["把某类任务的规则、提示、工具调用打包", "作用是快速做 demo，而不是保证结果"], (240, 230, 255)),
    ]
    for box, title, bullets, fill in boxes:
        draw_card(draw, box, title, bullets, fill)

    arrow_color = (116, 128, 150)
    draw.line((520, 288, 585, 288), fill=arrow_color, width=6)
    draw.line((1015, 288, 1080, 288), fill=arrow_color, width=6)
    draw.line((800, 365, 610, 480), fill=arrow_color, width=6)
    draw.line((1180, 365, 1060, 480), fill=arrow_color, width=6)
    draw.line((775, 580, 840, 580), fill=arrow_color, width=6)

    footer = "教学提示：普通人最容易卡住的不是不会问，而是不知道应该先学哪一层。"
    draw.text((70, 800), footer, font=body_font, fill=(80, 91, 108))
    image.save(path)
    return path


def build_diagram_content_flow() -> Path:
    path = IMG_DIR / "02_content_workflow.png"
    image = Image.new("RGB", (1600, 900), (247, 249, 252))
    draw = ImageDraw.Draw(image)
    title_font = get_font(46, bold=True)
    body_font = get_font(24)
    draw.text((70, 48), "内容创作者 AI Workflow", font=title_font, fill=(34, 44, 61))
    draw.text((70, 112), "核心原则：不是让 AI 一次写完，而是让它进入固定流程里的重复步骤。", font=body_font, fill=(97, 109, 126))

    steps = [
        ("1 选题", ["拆痛点", "找角度", "做选题池"], (255, 231, 214)),
        ("2 标题", ["多标题版本", "不同平台风格", "避免直接定终稿"], (255, 244, 214)),
        ("3 提纲", ["结构展开", "补观点", "列案例"], (225, 239, 255)),
        ("4 初稿", ["先由人定观点", "再让 AI 补表达", "保留个人判断"], (222, 244, 229)),
        ("5 改写", ["公众号版", "小红书版", "朋友圈版"], (240, 230, 255)),
        ("6 发布前", ["CTA 校准", "错别字检查", "封面与首图确认"], (232, 241, 236)),
    ]

    x = 80
    y = 230
    width = 220
    gap = 28
    for idx, (title, bullets, fill) in enumerate(steps):
        draw_card(draw, (x, y, x + width, y + 460), title, bullets, fill)
        if idx < len(steps) - 1:
            draw.line((x + width, y + 225, x + width + gap, y + 225), fill=(121, 134, 156), width=6)
        x += width + gap

    draw.text((70, 770), "训练营升级点：每一步都有人帮你纠偏，所以 workflow 才能从“会看”变成“会跑”。", font=body_font, fill=(80, 91, 108))
    image.save(path)
    return path


def build_diagram_funnel() -> Path:
    path = IMG_DIR / "03_private_domain_funnel.png"
    image = Image.new("RGB", (1600, 900), (250, 247, 243))
    draw = ImageDraw.Draw(image)
    title_font = get_font(46, bold=True)
    body_font = get_font(24)
    draw.text((70, 48), "私域转化漏斗图", font=title_font, fill=(34, 44, 61))
    draw.text((70, 112), "课程设计重点：公开资料负责吸引与筛选，训练营负责真正做出结果。", font=body_font, fill=(97, 109, 126))

    levels = [
        ("公开内容", ["短内容吸引", "建立认知", "评论与收藏"], (255, 231, 214), 1280),
        ("免费资料包", ["领取关键词", "加你进私域", "做第一轮筛选"], (255, 244, 214), 1040),
        ("低客单产品", ["模板包", "微课", "入门答疑"], (225, 239, 255), 800),
        ("系统课程", ["方法体系", "案例讲解", "标准流程"], (222, 244, 229), 560),
        ("训练营 / 陪跑", ["作业反馈", "实操纠偏", "结果交付"], (240, 230, 255), 340),
    ]

    y = 200
    center = 800
    height = 105
    for title, bullets, fill, width in levels:
        left = center - width // 2
        right = center + width // 2
        draw.rounded_rectangle((left, y, right, y + height), radius=20, fill=fill)
        draw.text((left + 26, y + 16), title, font=get_font(32, bold=True), fill=(24, 35, 52))
        bullet_text = " / ".join(bullets)
        draw.text((left + 28, y + 58), bullet_text, font=body_font, fill=(58, 72, 92))
        y += 125

    draw.text((70, 850), "讲课重点：demo 能让用户看到可能性；训练营才负责把可能性变成稳定输出。", font=body_font, fill=(80, 91, 108))
    image.save(path)
    return path


def build_diagram_learning_path() -> Path:
    path = IMG_DIR / "04_learning_paths.png"
    image = Image.new("RGB", (1600, 900), (244, 248, 246))
    draw = ImageDraw.Draw(image)
    title_font = get_font(46, bold=True)
    body_font = get_font(24)
    draw.text((70, 48), "三类人群学习路径图", font=title_font, fill=(34, 44, 61))
    draw.text((70, 112), "同一套课程体系，不同人群进入点不同，承接方式也不同。", font=body_font, fill=(97, 109, 126))

    columns = [
        ((80, 210, 500, 760), "新手", ["先搞懂概念", "先跑 1 条 workflow", "再决定学哪个工具"], (255, 244, 214)),
        ((590, 210, 1010, 760), "内容创作者", ["先搭内容 workflow", "先让 AI 提速", "再做多平台分发"], (225, 239, 255)),
        ((1100, 210, 1520, 760), "想变现的人", ["先做免费资料包", "先收私域", "再上低客单和训练营"], (222, 244, 229)),
    ]
    for box, title, bullets, fill in columns:
        draw_card(draw, box, title, bullets, fill)
        draw.text((box[0] + 28, box[3] - 120), "共同底层：workflow > 单个工具", font=get_font(22, bold=True), fill=(39, 53, 74))
        draw.text((box[0] + 28, box[3] - 82), "共同上层：训练营负责结果", font=get_font(22), fill=(58, 72, 92))

    image.save(path)
    return path


def set_run_font(run, name: str, size: int, bold: bool = False, color: tuple[int, int, int] = (0, 0, 0)) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(*color)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color in [
        ("Heading 1", 16, (46, 116, 181)),
        ("Heading 2", 13, (46, 116, 181)),
        ("Heading 3", 12, (31, 77, 120)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(*color)

    if "Card Title" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Card Title", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 77, 120)


def add_title_block(doc: Document, title: str, subtitle: str, tag: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_run_font(run, "Calibri", 24, bold=True, color=(34, 44, 61))

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(10)
    run2 = p2.add_run(subtitle)
    set_run_font(run2, "Calibri", 11, color=(90, 100, 120))

    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(16)
    run3 = p3.add_run(tag)
    set_run_font(run3, "Calibri", 10, bold=True, color=(210, 111, 48))


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run1 = p.add_run(bold_prefix)
        set_run_font(run1, "Calibri", 11, bold=True, color=(34, 44, 61))
        run2 = p.add_run(text[len(bold_prefix):])
        set_run_font(run2, "Calibri", 11, color=(34, 44, 61))
    else:
        run = p.add_run(text)
        set_run_font(run, "Calibri", 11, color=(34, 44, 61))


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, "Calibri", 11, color=(34, 44, 61))


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run, "Calibri", 11, color=(34, 44, 61))


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        shade_cell(cell, "E8EEF5")
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, "Calibri", 11, bold=True, color=(31, 77, 120))
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            p = cells[idx].paragraphs[0]
            run = p.add_run(value)
            set_run_font(run, "Calibri", 10, color=(34, 44, 61))
    doc.add_paragraph("")


def add_image(doc: Document, path: Path, title: str) -> None:
    p = doc.add_paragraph()
    p.style = "Card Title"
    p.add_run(title)
    doc.add_picture(str(path), width=Inches(6.2))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run("图示用于课程讲解与私域转化说明。")
    set_run_font(run, "Calibri", 9, color=(90, 100, 120))


def add_resource_section(doc: Document) -> None:
    doc.add_heading("公开学习资料与案例库", level=1)
    add_paragraph(doc, "以下资料适合做公开课、免费资料包、直播分享或课前阅读。重点不是全部学完，而是知道它们分别适合承担什么角色。")
    rows = [
        [resource.name, resource.category, resource.use_case, resource.note]
        for resource in PUBLIC_RESOURCES
    ]
    build_table(doc, ["名称", "类别", "适合怎么用", "课程里的讲法"], rows, [1.7, 1.3, 1.8, 1.7])
    add_paragraph(doc, "建议讲法：把这些公开资料定位为“快速出 demo 的公开底座”，而不是“学完就能出结果的完整训练”。")
    add_bullets(doc, [f"{resource.name}: {resource.url}" for resource in PUBLIC_RESOURCES])


def add_boundary_section(doc: Document) -> None:
    doc.add_heading("一定要讲透的边界", level=1)
    add_bullets(
        doc,
        [
            "公开资料和公开 skill 能帮助用户快速看懂、快速打样、快速上手。",
            "它们不保证内容质量、账号增长、转化结果和审美稳定性。",
            "真正的效果来自三件事：持续输入、具体反馈、重复练习。",
            "训练营存在的意义，不是多讲知识，而是帮用户把动作跑顺、把流程调对。",
        ],
    )


def build_total_handbook(diagrams: list[Path]) -> Path:
    doc = Document()
    configure_document(doc)
    add_title_block(
        doc,
        "AI 系统进阶课总手册",
        "面向新手、内容创作者与轻变现人群的公开资料版课程材料",
        "课程定位：公开资料负责认知与 demo，训练营负责效果与结果",
    )

    add_paragraph(
        doc,
        "这份总手册的目标，不是替代训练营，而是帮用户先看懂这门课到底在教什么、为什么要先学 workflow 而不是继续囤工具，以及不同人群该从哪里进入。",
    )

    doc.add_heading("课程总逻辑", level=1)
    add_numbered(
        doc,
        [
            "先建立正确认知：AI、大模型、agent、workflow、skill 各自是什么。",
            "再看公开 demo：哪些 GitHub skill 和 workflow 能快速出样。",
            "再做人群分流：新手、创作者、变现人群分别怎么学。",
            "最后把边界讲清楚：能看懂不等于能跑通，训练营负责从 demo 到结果。",
        ],
    )

    doc.add_heading("核心模块", level=1)
    for module in CORE_COURSE_MODULES:
        doc.add_heading(module["title"], level=2)
        add_bullets(doc, module["points"])

    for image_path, title in zip(
        diagrams,
        [
            "图 1：概念地图",
            "图 2：内容 workflow",
            "图 3：私域漏斗",
            "图 4：三类人群学习路径",
        ],
    ):
        add_image(doc, image_path, title)

    doc.add_heading("三类人群怎么承接", level=1)
    rows = []
    for persona in PERSONA_TRACKS.values():
        rows.append(
            [
                persona["title"].replace("给", "").replace("的", " "),
                " / ".join(persona["pain_points"][:2]),
                " / ".join(persona["goals"][:2]),
                persona["offer"],
            ]
        )
    build_table(doc, ["人群", "常见卡点", "先给什么结果", "后续承接"], rows, [1.2, 2.0, 2.0, 1.3])

    add_boundary_section(doc)
    add_resource_section(doc)

    doc.add_heading("建议的课程产品阶梯", level=1)
    add_bullets(
        doc,
        [
            "免费资料包：帮用户搞懂路线，顺便筛掉只想白嫖、不想行动的人。",
            "9.9-49.9 模板包或清单包：让用户感觉自己开始能动起来。",
            "69-199 微课：把一条 workflow 讲顺。",
            "299-999 系统课：把认知、流程、工具和案例串起来。",
            "999+ 训练营：给作业、改动作、做结果。",
        ],
    )

    path = OUT_DIR / "AI系统进阶课-总手册.docx"
    doc.save(path)
    return path


def build_persona_doc(key: str, diagrams: list[Path]) -> Path:
    persona = PERSONA_TRACKS[key]
    doc = Document()
    configure_document(doc)
    add_title_block(doc, persona["title"], persona["subtitle"], persona["positioning"])

    doc.add_heading("这份手册适合谁", level=1)
    add_bullets(doc, persona["pain_points"])

    doc.add_heading("学完以后先拿到什么", level=1)
    add_bullets(doc, persona["goals"])

    doc.add_heading("建议的学习顺序", level=1)
    add_numbered(doc, persona["modules"])

    doc.add_heading("你需要建立的底层认知", level=1)
    add_bullets(
        doc,
        [
            "工具不是重点，流程才是重点。",
            "大模型提供能力，agent 提供行动，workflow 提供稳定，skill 提供复用。",
            "公开 skill 最适合做 demo 和入门，不适合承诺最终效果。",
        ],
    )

    if key == "newbie":
        add_image(doc, diagrams[0], "概念地图")
        doc.add_heading("新手 14 天入门路线", level=1)
        add_numbered(
            doc,
            [
                "第 1-3 天：理解 AI、大模型、agent、workflow、skill 的区别。",
                "第 4-6 天：学会把一个任务拆成 4 到 6 步，而不是一句话全交给 AI。",
                "第 7-10 天：用公开资料和公开 skill 跑一条最简单的内容或整理 demo。",
                "第 11-14 天：复盘哪里卡住，决定后面是走内容提效路线还是轻变现路线。",
            ],
        )
    elif key == "creator":
        add_image(doc, diagrams[1], "内容 workflow")
        doc.add_heading("内容创作者最该先学的一条流程", level=1)
        add_bullets(
            doc,
            [
                "先拆受众痛点，再列内容角度。",
                "先定观点和结构，再让 AI 帮你补表达、补例子、补转场。",
                "最后再改成平台版本，而不是一开始就让 AI 直接写终稿。",
            ],
        )
    elif key == "monetizer":
        add_image(doc, diagrams[2], "私域漏斗")
        doc.add_heading("轻变现人群先跑哪条闭环", level=1)
        add_numbered(
            doc,
            [
                "先用短内容吸引注意力。",
                "再用免费资料包完成加你和筛选。",
                "再用低客单产品验证付费意愿。",
                "最后用系统课和训练营交付结果。",
            ],
        )

    doc.add_heading("公开 skill 和 workflow 怎么讲", level=1)
    rows = [
        [resource.name, resource.use_case, resource.note]
        for resource in PUBLIC_RESOURCES
        if "GitHub" in resource.category
    ]
    build_table(doc, ["公开仓库", "适合演示什么", "不要夸大的地方"], rows, [2.0, 2.0, 2.5])

    add_boundary_section(doc)
    doc.add_heading("最适合的后续承接", level=1)
    add_paragraph(doc, persona["offer"])
    add_paragraph(
        doc,
        "建议讲法：公开资料先让用户看到“自己可能也能做到”，训练营再让用户真正做到“我已经把它跑通了”。",
    )

    path = OUT_DIR / f"{persona['title']}.docx"
    doc.save(path)
    return path


def build_contact_sheet(source_paths: list[Path], output_path: Path) -> None:
    images = [Image.open(path).convert("RGB") for path in source_paths]
    thumb_width = 320
    thumbs = []
    for img in images:
        ratio = thumb_width / img.width
        size = (thumb_width, int(img.height * ratio))
        thumbs.append(img.resize(size))
    cols = 2
    rows = (len(thumbs) + cols - 1) // cols
    padding = 30
    canvas_width = cols * thumb_width + (cols + 1) * padding
    row_heights = []
    for row in range(rows):
        row_imgs = thumbs[row * cols:(row + 1) * cols]
        row_heights.append(max(img.height for img in row_imgs))
    canvas_height = sum(row_heights) + (rows + 1) * padding
    canvas = Image.new("RGB", (canvas_width, canvas_height), (248, 248, 248))
    y = padding
    idx = 0
    label_font = get_font(18, bold=True)
    for row, row_height in enumerate(row_heights):
        x = padding
        for _ in range(cols):
            if idx >= len(thumbs):
                break
            canvas.paste(thumbs[idx], (x, y))
            label = source_paths[idx].stem
            ImageDraw.Draw(canvas).text((x, y + thumbs[idx].height + 4), label, font=label_font, fill=(50, 60, 75))
            x += thumb_width + padding
            idx += 1
        y += row_height + 38
    canvas.save(output_path)


def main() -> None:
    ensure_dirs()
    diagrams = [
        build_diagram_concept_map(),
        build_diagram_content_flow(),
        build_diagram_funnel(),
        build_diagram_learning_path(),
    ]
    docs = [build_total_handbook(diagrams)]
    docs.append(build_persona_doc("newbie", diagrams))
    docs.append(build_persona_doc("creator", diagrams))
    docs.append(build_persona_doc("monetizer", diagrams))

    with open(OUT_DIR / "README.txt", "w", encoding="utf-8") as fh:
        fh.write("已生成课程材料：\n")
        for doc in docs:
            fh.write(f"- {doc.name}\n")
        fh.write("\n配图：\n")
        for img in diagrams:
            fh.write(f"- {img.name}\n")

    with open(OUT_DIR / "公开资料与案例索引.md", "w", encoding="utf-8") as fh:
        fh.write("# 公开资料与案例索引\n\n")
        fh.write("这份索引用于后续直播、海报、资料包、私域答疑和课程延伸阅读。\n\n")
        for resource in PUBLIC_RESOURCES:
            fh.write(f"## {resource.name}\n")
            fh.write(f"- 类别：{resource.category}\n")
            fh.write(f"- 地址：{resource.url}\n")
            fh.write(f"- 适合用途：{resource.use_case}\n")
            fh.write(f"- 课程里的讲法：{resource.note}\n\n")


if __name__ == "__main__":
    main()
