from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_course_materials import (
    PUBLIC_RESOURCES,
    ROOT,
    add_bullets,
    add_image,
    add_numbered,
    add_paragraph,
    add_title_block,
    build_diagram_concept_map,
    build_diagram_content_flow,
    build_diagram_funnel,
    build_diagram_learning_path,
    build_table,
    configure_document,
    ensure_dirs,
    get_font,
    IMG_DIR,
    wrap_text,
)


OUT_DIR = ROOT / "deliverables" / "course_library_v2"
CATEGORY_IMG_DIR = OUT_DIR / "_category_images"


CATEGORIES = [
    {
        "folder": "01_AI基础认知",
        "title": "AI 基础认知资料库",
        "intro": "先把概念讲清楚，再谈工具、项目和变现。",
        "docs": [
            {
                "name": "01_什么是AI.docx",
                "title": "什么是 AI",
                "subtitle": "给第一次系统接触 AI 的用户",
                "goal": "先建立最基本的概念边界，避免把任何自动化都叫 AI。",
                "for_whom": ["第一次系统学 AI 的人", "准备做 AI 公开课的人", "想给私域用户讲清楚概念的人"],
                "sections": [
                    ("先讲结论", ["AI 不是某一个工具，而是一类让机器完成识别、生成、预测和辅助判断任务的方法。", "你在课程里不需要把 AI 讲成高深技术，更需要讲清楚它解决什么问题。"]),
                    ("课程里最该强调什么", ["不要把 AI 讲成万能替代者。", "把 AI 讲成‘提高速度、扩展选项、降低重复劳动’更容易让普通人理解。", "AI 的价值不在于炫技，而在于是否进入了具体业务步骤。"]),
                    ("用户最容易误解的地方", ["把 AI 当作会自己负责结果的人。", "把会聊天和会做事混为一谈。", "看太多案例，却没有自己的任务场景。"]),
                ],
                "actions": ["用一句话解释 AI 是什么。", "列出你自己工作里最重复的 3 个动作。", "判断这 3 个动作里哪些适合让 AI 协助。"],
                "refs": ["IBM LLM overview", "OpenAI Cookbook"],
                "image": "concept",
            },
            {
                "name": "02_什么是生成式AI与大模型.docx",
                "title": "什么是生成式 AI 与大模型",
                "subtitle": "把最常混淆的两个词拆开讲",
                "goal": "帮助用户理解大模型只是底层能力，生成式 AI 是应用方向。",
                "for_whom": ["被各种 AI 名词绕晕的人", "准备做 AI 入门资料包的人"],
                "sections": [
                    ("概念拆分", ["生成式 AI 侧重‘能生成内容’。", "大模型是支撑很多生成式 AI 产品的底层通用模型能力。", "所以不是所有 AI 都是大模型，也不是所有大模型只做聊天。"]),
                    ("大模型为什么强", ["它能处理语言、图像甚至多模态输入。", "它擅长从大量数据中学习模式，再按提示生成结果。"]),
                    ("课程里怎么讲更接地气", ["把大模型比喻成通用发动机。", "把应用产品比喻成装在发动机上的不同车身。"]),
                ],
                "actions": ["区分‘模型’和‘产品’。", "列出你正在用的 3 个 AI 产品，并写出它们背后依赖的是哪类模型能力。"],
                "refs": ["IBM LLM overview", "OpenAI Models and prompt guides"],
                "image": "concept",
            },
            {
                "name": "03_什么是Prompt与提示工程.docx",
                "title": "什么是 Prompt 与提示工程",
                "subtitle": "先学会表达任务，再学更复杂的 agent",
                "goal": "让新手理解 prompt 不是魔法咒语，而是任务说明书。",
                "for_whom": ["一上来就搜提示词模板的人", "想把提示词做成资料的人"],
                "sections": [
                    ("先讲定义", ["Prompt 是你给模型的任务输入。", "提示工程不是追求华丽词藻，而是让任务、上下文、输出格式更清晰。"]),
                    ("好 prompt 的共同特征", ["任务目标明确。", "上下文充分。", "输出格式清楚。", "边界条件明确。"]),
                    ("最常见的 3 个坑", ["一句话把所有事都丢给模型。", "没有提供背景和对象。", "没有说清楚最终要什么格式。"]),
                ],
                "actions": ["把一句模糊问题改成结构化提示。", "给自己的一个常见任务写出目标、背景、格式三段。"],
                "refs": ["OpenAI Prompt engineering", "OpenAI Cookbook"],
                "image": None,
            },
            {
                "name": "04_什么是Agent.docx",
                "title": "什么是 Agent",
                "subtitle": "从会聊天到会行动，中间多了哪些能力",
                "goal": "让用户理解 agent 的价值在于分步执行和调用工具。",
                "for_whom": ["听过 agent 但不知道和普通聊天有什么区别的人", "准备讲 agent 公开课的人"],
                "sections": [
                    ("一句话理解", ["Agent 是能理解目标、调用工具、按步骤执行并根据结果继续行动的 AI 形态。"]),
                    ("它和普通聊天的差别", ["普通聊天偏回答问题。", "Agent 更偏完成任务。", "一旦涉及搜索、调用接口、写文件、操作浏览器，agent 的价值会明显提高。"]),
                    ("为什么很多人学了还不会用", ["因为没有实际任务。", "因为不会拆步骤。", "因为没有把‘判断’和‘执行’分开。"]),
                ],
                "actions": ["找一件 10 到 20 分钟的重复任务，把它拆成 5 步。", "标出其中哪些步可以交给 agent。"],
                "refs": ["OpenAI Agents SDK", "oTTomator Agents"],
                "image": "concept",
            },
            {
                "name": "05_什么是Workflow与Skill.docx",
                "title": "什么是 Workflow 与 Skill",
                "subtitle": "为什么真正值钱的是流程，而不是收藏更多工具",
                "goal": "把 workflow 和 skill 的职责边界讲清楚。",
                "for_whom": ["一直在找现成 skill，但没有自己的流程的人", "想卖资料包和流程包的人"],
                "sections": [
                    ("核心区别", ["Workflow 是任务拆解顺序。", "Skill 是把某一类动作打包成可复用能力。"]),
                    ("为什么课程里要先讲 workflow", ["因为先有任务步骤，才知道 skill 装在哪一步。", "没有流程，skill 只会变成又一个收藏夹。"]),
                    ("为什么 skill 适合做 demo", ["上手快。", "可以快速让用户看到结果样子。", "但不自动保证结果稳定。"]),
                ],
                "actions": ["写出你的一条 workflow。", "从 workflow 中找出最适合封装成 skill 的 2 步。"],
                "refs": ["OpenAI Agents and tools docs", "blacktwist/social-media-skills"],
                "image": "concept",
            },
            {
                "name": "06_AI学习常见误区.docx",
                "title": "AI 学习常见误区",
                "subtitle": "这份文档适合直接拿去做公开课答疑",
                "goal": "降低用户对 AI 学习路径的误判。",
                "for_whom": ["容易被热点带节奏的人", "你私域里的新手用户"],
                "sections": [
                    ("误区 1", ["学更多工具，不等于结果更好。"]),
                    ("误区 2", ["知道很多案例，不等于自己会跑。"]),
                    ("误区 3", ["有了 prompt 模板，不等于有了 workflow。"]),
                    ("误区 4", ["看到 demo，不等于拿到结果。训练营的价值就在这里。"]),
                ],
                "actions": ["删掉暂时不会用到的工具收藏。", "选一条最短 workflow 先跑通。", "在实际任务里连续用 7 天再判断值不值得学。"],
                "refs": ["IBM LLM overview", "OpenAI Prompt guides", "公开 GitHub skills examples"],
                "image": "learning",
            },
        ],
    },
    {
        "folder": "02_提示词与模型使用",
        "title": "提示词与模型使用资料库",
        "intro": "这一组文档偏实操，适合做模板包、微课和直播分享。",
        "docs": [
            {
                "name": "01_新手提问法.docx",
                "title": "新手提问法",
                "subtitle": "不会问，后面所有 skill 和 workflow 都很难稳定",
                "goal": "先把模糊需求改造成可执行任务。",
                "for_whom": ["问一句‘帮我写一个’的人", "需要提升模型配合度的人"],
                "sections": [
                    ("最小提问结构", ["目标是什么。", "对象是谁。", "你已经有什么。", "最后想要什么格式。"]),
                    ("3 种提问层次", ["开放型提问：适合找角度。", "约束型提问：适合出结果。", "迭代型提问：适合修稿和纠偏。"]),
                    ("教学建议", ["别先讲复杂框架，先把‘任务-背景-输出’三件事讲熟。"]),
                ],
                "actions": ["把你最近的一个提问，重写成四段结构。"],
                "refs": ["OpenAI Prompt engineering", "OpenAI Cookbook"],
                "image": None,
            },
            {
                "name": "02_结构化提示模板.docx",
                "title": "结构化提示模板",
                "subtitle": "适合整理成你的私域免费资料",
                "goal": "给用户一套通用的提示骨架，而不是零碎句子。",
                "for_whom": ["做模板包的人", "经常重复写同类任务的人"],
                "sections": [
                    ("通用模板", ["角色", "任务", "背景", "限制", "输出格式", "示例标准"]),
                    ("什么时候一定要结构化", ["内容写作。", "信息整理。", "报告输出。", "多步骤工作。"]),
                    ("什么时候不用过度复杂", ["只是要几个灵感角度时。", "只是要快速头脑风暴时。"]),
                ],
                "actions": ["给选题、写作、总结各写一个结构化模板。"],
                "refs": ["OpenAI Prompt engineering", "OpenAI Cookbook"],
                "image": None,
            },
            {
                "name": "03_长文写作提示.docx",
                "title": "长文写作提示",
                "subtitle": "把大模型从写稿器变成写作协作器",
                "goal": "教会用户写长文时如何拆成多轮，而不是一轮出终稿。",
                "for_whom": ["写公众号、课程讲义、长文案的人"],
                "sections": [
                    ("先拆 4 段", ["目标读者。", "文章观点。", "结构提纲。", "段落扩写。"]),
                    ("为什么不能直接让 AI 写终稿", ["平均、安全、像模板。", "缺少真实观点。", "难保留个人表达。"]),
                    ("更稳的做法", ["先由人定观点。", "再让 AI 补结构和表达。", "最后自己做风格修正。"]),
                ],
                "actions": ["把一篇长文拆成提纲 prompt、扩写 prompt、改写 prompt 三段。"],
                "refs": ["OpenAI Prompt engineering", "OpenAI Cookbook"],
                "image": None,
            },
            {
                "name": "04_信息整理和总结提示.docx",
                "title": "信息整理和总结提示",
                "subtitle": "这是普通人最容易先拿到提效感的一类任务",
                "goal": "先从整理类任务入手，让用户更容易体验到 AI 价值。",
                "for_whom": ["做资料整理、会议纪要、知识汇总的人"],
                "sections": [
                    ("适合 AI 的整理任务", ["长文摘要。", "评论归类。", "会议纪要提炼。", "课程笔记整理。"]),
                    ("输出格式怎么定", ["摘要版。", "表格版。", "清单版。", "待办版。"]),
                    ("注意事项", ["原文越乱，越要先说明整理标准。", "敏感信息要先脱敏。"]),
                ],
                "actions": ["用一个真实素材练 3 种输出格式。"],
                "refs": ["OpenAI Cookbook", "OpenAI Prompt guides"],
                "image": None,
            },
            {
                "name": "05_角色提示与风格控制.docx",
                "title": "角色提示与风格控制",
                "subtitle": "减少 AI 味，重点不在修辞，而在上下文和标准",
                "goal": "讲清楚角色提示的作用边界。",
                "for_whom": ["内容创作者", "想让输出更像自己的人"],
                "sections": [
                    ("角色提示能解决什么", ["帮助模型理解任务视角。", "帮助稳定输出结构。"]),
                    ("角色提示解决不了什么", ["不能凭空创造真实经历。", "不能替代你的观点和审美。"]),
                    ("风格控制更关键的部分", ["给样例。", "给禁忌。", "给语气边界。"]),
                ],
                "actions": ["为你自己的账号写一版角色设定。", "写 3 条必须避免的表达习惯。"],
                "refs": ["OpenAI Prompt engineering", "blacktwist/social-media-skills"],
                "image": None,
            },
            {
                "name": "06_提示迭代与复盘.docx",
                "title": "提示迭代与复盘",
                "subtitle": "真正有用的不是一个神 prompt，而是可复盘的改法",
                "goal": "把提示词学习从收藏转成迭代。",
                "for_whom": ["积累很多 prompt 却不知道怎么优化的人"],
                "sections": [
                    ("复盘看什么", ["目标是否清楚。", "背景是否足够。", "格式是否明确。", "结果差在哪一步。"]),
                    ("3 种常见改法", ["补上下文。", "缩小任务范围。", "先分步后合并。"]),
                    ("适合做资料包的形式", ["错例对比。", "前后版本对比。", "模板升级记录。"]),
                ],
                "actions": ["把一个失败 prompt 改 3 版，并记录变化。"],
                "refs": ["OpenAI Prompt engineering", "OpenAI Cookbook"],
                "image": None,
            },
        ],
    },
    {
        "folder": "03_Agent_Skills_Workflows",
        "title": "Agent / Skills / Workflows 资料库",
        "intro": "这一组适合做系统课和训练营的中间层，把‘会问’升级成‘会跑流程’。",
        "docs": [
            {
                "name": "01_Agent基本结构.docx",
                "title": "Agent 基本结构",
                "subtitle": "理解 agent 时，重点是目标、工具、状态和判断",
                "goal": "让用户看懂 agent 不只是聊天机器人升级版。",
                "for_whom": ["准备做 agent 入门课的人", "想从 prompt 升级到自动化的人"],
                "sections": [
                    ("最小结构", ["目标。", "模型。", "工具。", "状态。", "输出。"]),
                    ("为什么会比单次提问更强", ["可以分步执行。", "可以根据中间结果继续行动。", "可以调用外部工具。"]),
                    ("课程里怎么讲", ["把 agent 讲成‘会调工具的执行者’最容易理解。"]),
                ],
                "actions": ["给一个 agent 任务写出目标、工具和成功标准。"],
                "refs": ["OpenAI Agents SDK", "oTTomator Agents"],
                "image": "concept",
            },
            {
                "name": "02_Tools函数调用和MCP概念.docx",
                "title": "Tools、函数调用和 MCP 概念",
                "subtitle": "为什么 agent 能真正做事，关键在工具层",
                "goal": "让用户理解函数调用、外部工具和 MCP 的关系。",
                "for_whom": ["开始接触 agent 架构的人", "准备解释工具调用的人"],
                "sections": [
                    ("函数调用", ["模型输出结构化参数，让程序去执行具体动作。"]),
                    ("工具层", ["搜索、文件、浏览器、数据库、接口，都是 agent 的手和脚。"]),
                    ("MCP 的课程讲法", ["可以把 MCP 讲成统一接插件接口，帮助模型安全地接入工具和数据。"]),
                ],
                "actions": ["列出你最想接入 agent 的 5 类工具。"],
                "refs": ["OpenAI tools docs", "OpenAI agents docs"],
                "image": None,
            },
            {
                "name": "03_Skills如何快速出Demo.docx",
                "title": "Skills 如何快速出 Demo",
                "subtitle": "skill 的作用是打包经验，不是直接替你拿结果",
                "goal": "讲清楚 skill 适合做课程演示、模板打样和流程封装。",
                "for_whom": ["想卖 skill 资料包的人", "想给用户展示 AI 能做什么的人"],
                "sections": [
                    ("Skill 的价值", ["统一输入。", "统一输出。", "统一操作顺序。"]),
                    ("适合做什么 demo", ["写作。", "改写。", "发布。", "搜索。", "内容分析。"]),
                    ("不该夸大的地方", ["skill 不会自动知道你的业务判断。", "skill 不会天然带来流量和成交。"]),
                ],
                "actions": ["把一个重复任务写成简化版 skill 说明。"],
                "refs": ["blacktwist/social-media-skills", "xiaohongshu-mcp-skills"],
                "image": "concept",
            },
            {
                "name": "04_Workflow如何拆步骤.docx",
                "title": "Workflow 如何拆步骤",
                "subtitle": "一切稳定产出都从任务拆解开始",
                "goal": "让用户学会把复杂任务拆成稳定动作。",
                "for_whom": ["总觉得 AI 不稳定的人", "想做 SOP 和训练营的人"],
                "sections": [
                    ("拆流程的 4 个标准", ["先后顺序。", "谁负责判断。", "哪里适合 AI。", "哪里需要人工确认。"]),
                    ("为什么很多 demo 没法复用", ["没有写清楚中间步骤。", "每次都靠临场发挥。"]),
                    ("课程里怎么教", ["先画流程图，再补工具。"]),
                ],
                "actions": ["画出你当前业务里最重要的一条 6 步流程。"],
                "refs": ["OpenAI agents orchestration docs", "公开 GitHub workflow repos"],
                "image": "content",
            },
            {
                "name": "05_浏览器自动化与发布Demo.docx",
                "title": "浏览器自动化与发布 Demo",
                "subtitle": "从会写文案，到会执行动作，中间需要浏览器层能力",
                "goal": "让用户理解浏览器自动化适合做半自动执行，不适合夸大成全自动生意系统。",
                "for_whom": ["想讲自动发布、自动操作的人", "做社媒流程的人"],
                "sections": [
                    ("适合演示什么", ["登录后的页面操作。", "草稿填写。", "表单提交。", "批量重复动作。"]),
                    ("为什么适合做 demo", ["视觉冲击强。", "用户容易感知价值。"]),
                    ("边界", ["平台规则会变。", "账号安全要考虑。", "关键动作最好保留人工确认。"]),
                ],
                "actions": ["找 1 个需要重复点点点的动作，写成半自动 demo 方案。"],
                "refs": ["agent-browser", "social-push"],
                "image": None,
            },
            {
                "name": "06_公开Agent项目案例库.docx",
                "title": "公开 Agent 项目案例库",
                "subtitle": "这份适合拿来做课程案例、直播演示和灵感扩展",
                "goal": "给用户看更广的 agent 用例，打开想象空间。",
                "for_whom": ["准备扩展课程深度的人", "想给用户更多案例的人"],
                "sections": [
                    ("推荐案例方向", ["内容创作 agent。", "研究 agent。", "RAG agent。", "浏览器操作 agent。", "线索收集 agent。"]),
                    ("为什么要看公开项目", ["能看见不同结构。", "能理解 agent 不止一种写法。"]),
                    ("课程里的讲法", ["不是让用户全部学会，而是帮助他们知道‘未来能做到什么’。"]),
                ],
                "actions": ["挑 3 个公开 agent 项目，写出它们分别解决什么问题。"],
                "refs": ["oTTomator Agents", "OpenAI Cookbook"],
                "image": None,
            },
        ],
    },
    {
        "folder": "04_内容创作实战",
        "title": "内容创作实战资料库",
        "intro": "这一组直接服务内容创作者和自媒体用户，也是最容易承接到私域的方向。",
        "docs": [
            {
                "name": "01_内容选题Workflow.docx",
                "title": "内容选题 Workflow",
                "subtitle": "不要先写，先拆用户和问题",
                "goal": "把内容选题从灵感型，变成流程型。",
                "for_whom": ["做公众号、小红书、私域内容的人"],
                "sections": [
                    ("选题先看什么", ["人群。", "痛点。", "场景。", "转化动作。"]),
                    ("AI 适合帮哪几步", ["拆问题。", "列角度。", "补反问。", "整理选题池。"]),
                    ("最常见问题", ["标题先行。", "热点先行。", "没有承接目标。"]),
                ],
                "actions": ["围绕一类人群写出 10 个内容角度。"],
                "refs": ["blacktwist/social-media-skills", "OpenAI Prompt engineering"],
                "image": "content",
            },
            {
                "name": "02_标题与钩子.docx",
                "title": "标题与钩子",
                "subtitle": "标题不是文学创作，而是筛选点击和人群",
                "goal": "讲清楚标题和开头的任务不同，但都要服务转化。",
                "for_whom": ["短内容创作者", "需要提升打开率的人"],
                "sections": [
                    ("标题的任务", ["让对的人停下来。", "让用户知道这条内容跟自己有关。"]),
                    ("钩子的任务", ["在前 1 到 3 句交代冲突、结果或反差。"]),
                    ("AI 适合怎么协助", ["给多个角度。", "做语气切换。", "改成平台版本。"]),
                ],
                "actions": ["用同一主题写 5 个不同风格标题。"],
                "refs": ["blacktwist/social-media-skills", "公开社媒技能 README"],
                "image": None,
            },
            {
                "name": "03_提纲与初稿协作.docx",
                "title": "提纲与初稿协作",
                "subtitle": "先由人定结构，再让 AI 补表达",
                "goal": "减少 AI 味，提升你对内容骨架的掌控。",
                "for_whom": ["写长图文、长文案、课程文稿的人"],
                "sections": [
                    ("人先做什么", ["定观点。", "定顺序。", "定重点。"]),
                    ("AI 再做什么", ["扩写。", "补例子。", "补转场。", "给不同说法。"]),
                    ("最关键的一点", ["不要让 AI 替你决定内容立场。"]),
                ],
                "actions": ["用同一个提纲，让 AI 分别做扩写和改写。"],
                "refs": ["OpenAI Prompt engineering", "OpenAI Cookbook"],
                "image": "content",
            },
            {
                "name": "04_改写与多平台复用.docx",
                "title": "改写与多平台复用",
                "subtitle": "一份内容，多次产出，关键不是复制，而是重写",
                "goal": "帮助用户把同一主题改成公众号、小红书、朋友圈等版本。",
                "for_whom": ["多平台运营者", "内容量不足的人"],
                "sections": [
                    ("为什么不能直接复制", ["平台语气不同。", "平台长度不同。", "平台用户期待不同。"]),
                    ("改写时最该保留什么", ["核心观点。", "用户痛点。", "最终 CTA。"]),
                    ("AI 最适合做哪一步", ["改语气。", "改长短。", "改结构。"]),
                ],
                "actions": ["把一篇长文改成小红书图文版和朋友圈短帖版。"],
                "refs": ["blacktwist/social-media-skills", "social-push"],
                "image": None,
            },
            {
                "name": "05_配图与素材整理.docx",
                "title": "配图与素材整理",
                "subtitle": "视觉素材不是附属品，而是内容完成度的一部分",
                "goal": "让用户理解图片、封面、首图和素材整理也可以进入 workflow。",
                "for_whom": ["图文内容创作者", "想系统做资料包的人"],
                "sections": [
                    ("该整理哪些素材", ["首图模板。", "案例截图。", "流程图。", "清单卡片。"]),
                    ("AI 能帮什么", ["给图示结构。", "生成文案说明。", "整理视觉说明书。"]),
                    ("边界", ["审美和品牌统一仍然要人把关。"]),
                ],
                "actions": ["为一套资料包整理首图、流程图、要点图的素材清单。"],
                "refs": ["OpenAI Cookbook", "公开视觉教程和平台规则"],
                "image": "learning",
            },
            {
                "name": "06_发布前检查与数据复盘.docx",
                "title": "发布前检查与数据复盘",
                "subtitle": "提效不只是生成，还包括检查和复盘",
                "goal": "把发布前检查和发布后复盘纳入内容 workflow。",
                "for_whom": ["内容创作者", "团队运营人员"],
                "sections": [
                    ("发布前检查什么", ["标题是否匹配内容。", "CTA 是否单一。", "错别字与口径是否一致。"]),
                    ("发布后复盘什么", ["收藏评论。", "私信关键词。", "加你比例。", "成交相关反馈。"]),
                    ("为什么这一步重要", ["没有复盘，后面所有 skill 都不会进化。"]),
                ],
                "actions": ["为你自己的内容制作一张发布前后检查表。"],
                "refs": ["blacktwist/social-media-skills", "social-push"],
                "image": "funnel",
            },
        ],
    },
    {
        "folder": "05_私域转化与训练营",
        "title": "私域转化与训练营资料库",
        "intro": "这一组直接服务变现和承接，把公开资料变成产品路径。",
        "docs": [
            {
                "name": "01_免费资料包设计.docx",
                "title": "免费资料包设计",
                "subtitle": "资料包不是白送内容，而是筛选器和入口",
                "goal": "让用户理解免费资料的真正任务是引流和筛选。",
                "for_whom": ["想做私域承接的人", "准备卖课或训练营的人"],
                "sections": [
                    ("资料包最该解决什么", ["让用户更清楚下一步。", "让用户愿意加你。", "帮你判断对方属于哪类人。"]),
                    ("适合做成资料包的主题", ["5 条 workflow。", "10 个常见误区。", "1 套内容 SOP。", "1 套模板包。"]),
                    ("不建议做什么", ["大而全百科。", "看完就没有下一步的资料。"]),
                ],
                "actions": ["写出你的资料包题目、对象、CTA 和后续承接。"],
                "refs": ["公开课程营销案例", "你当前课程体系整理"],
                "image": "funnel",
            },
            {
                "name": "02_留资与加微信路径.docx",
                "title": "留资与加微信路径",
                "subtitle": "内容不是终点，进入私域才是承接开始",
                "goal": "把留资动作设计得更清楚、更自然。",
                "for_whom": ["做小红书、公众号、社群引流的人"],
                "sections": [
                    ("常见留资动作", ["评论关键词。", "私信关键词。", "公众号回复。", "表单登记。"]),
                    ("什么样的 CTA 更容易转化", ["明确。", "单一。", "低门槛。", "对应内容主题。"]),
                    ("不要怎么做", ["一条内容放多个 CTA。", "内容和资料包主题不对应。"]),
                ],
                "actions": ["为 3 条内容分别写一个单一 CTA。"],
                "refs": ["公开社媒自动化 workflow", "平台内容策略资料"],
                "image": "funnel",
            },
            {
                "name": "03_低客单产品设计.docx",
                "title": "低客单产品设计",
                "subtitle": "低客单不是为了赚最多，而是为了验证和筛人",
                "goal": "帮助你设计 9.9-199 区间的产品承接层。",
                "for_whom": ["想做第一笔成交的人", "资料包后续承接不足的人"],
                "sections": [
                    ("低客单适合卖什么", ["模板包。", "清单包。", "微课。", "单点答疑。"]),
                    ("低客单的真正作用", ["验证付费意愿。", "筛选高意向用户。", "为系统课和训练营铺路。"]),
                    ("课程里最该讲的逻辑", ["不是卖便宜，而是卖确定的下一步。"]),
                ],
                "actions": ["设计一个 49 元或 99 元的过渡产品。"],
                "refs": ["公开课程产品案例", "私域转化常见结构"],
                "image": "funnel",
            },
            {
                "name": "04_系统课结构设计.docx",
                "title": "系统课结构设计",
                "subtitle": "把碎片知识整理成递进式课程结构",
                "goal": "帮助你把公开资料升级成系统课程框架。",
                "for_whom": ["准备做 299-999 系统课的人"],
                "sections": [
                    ("系统课应该怎么分层", ["认知层。", "工具层。", "workflow 层。", "案例层。", "转化层。"]),
                    ("为什么别一上来做太大", ["没有验证需求。", "没有真实问题池。", "用户也消化不了。"]),
                    ("更稳的做法", ["先从资料包和低客单中收需求，再整理成系统课。"]),
                ],
                "actions": ["用 5 个模块写出你的系统课大纲。"],
                "refs": ["公开课程目录设计方法", "当前资料库结构"],
                "image": "learning",
            },
            {
                "name": "05_训练营设计与作业机制.docx",
                "title": "训练营设计与作业机制",
                "subtitle": "训练营的价值不在知识量，而在动作矫正和结果交付",
                "goal": "讲清楚为什么你必须把训练营设计成反馈系统。",
                "for_whom": ["想做陪跑营、训练营的人"],
                "sections": [
                    ("训练营最该交付什么", ["动作改变。", "流程跑通。", "第一批结果。"]),
                    ("为什么公开资料做不到", ["没有作业。", "没有反馈。", "没有人监督。"]),
                    ("训练营设计最少要有", ["固定作业。", "案例点评。", "群内答疑。", "结果节点。"]),
                ],
                "actions": ["为你的训练营写出 7 天或 14 天作业安排。"],
                "refs": ["公开训练营设计方法", "私域陪跑案例"],
                "image": "funnel",
            },
            {
                "name": "06_常见成交问题答疑.docx",
                "title": "常见成交问题答疑",
                "subtitle": "这份文档适合直接拿去做私域 FAQ",
                "goal": "提前准备用户最常问的问题，降低成交阻力。",
                "for_whom": ["私域运营者", "课程销售转化承接者"],
                "sections": [
                    ("常见问题 1", ["学完是不是马上见效。", "回答重点：公开资料看懂路线，训练营才负责做出结果。"]),
                    ("常见问题 2", ["我是不是零基础也能学。", "回答重点：可以，但先从认知和简单 workflow 开始。"]),
                    ("常见问题 3", ["为什么不能只买资料。", "回答重点：资料能帮你知道怎么做，但不会替你纠偏。"]),
                ],
                "actions": ["把你私域里最常见的 10 个问题整理成标准答法。"],
                "refs": ["公开课程 FAQ 结构", "你自己的私域承接逻辑"],
                "image": None,
            },
        ],
    },
]


IMAGE_MAP = {
    "concept": "01_ai_concept_map.png",
    "content": "02_content_workflow.png",
    "funnel": "03_private_domain_funnel.png",
    "learning": "04_learning_paths.png",
}


def ensure_output() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    CATEGORY_IMG_DIR.mkdir(parents=True, exist_ok=True)
    for category in CATEGORIES:
        (OUT_DIR / category["folder"]).mkdir(parents=True, exist_ok=True)


def image_path(name: str | None) -> Path | None:
    if not name:
        return None
    return IMG_DIR / IMAGE_MAP[name]


def build_category_banner(category: dict) -> Path:
    path = CATEGORY_IMG_DIR / f"{category['folder']}.png"
    image = Image.new("RGB", (1600, 900), (247, 245, 240))
    draw = ImageDraw.Draw(image)
    title_font = get_font(52, bold=True)
    sub_font = get_font(26)
    body_font = get_font(28)

    draw.text((70, 60), category["title"], font=title_font, fill=(35, 45, 63))
    lines = wrap_text(draw, category["intro"], sub_font, 1400)
    y = 145
    for line in lines:
        draw.text((70, y), line, font=sub_font, fill=(93, 104, 122))
        y += 40

    colors = [
        (255, 232, 214),
        (252, 241, 206),
        (221, 234, 251),
        (220, 239, 229),
        (232, 224, 248),
        (229, 237, 236),
    ]
    x = 70
    y = 280
    for idx, item in enumerate(category["docs"][:6]):
        fill = colors[idx % len(colors)]
        box = (x, y, x + 460, y + 190)
        draw.rounded_rectangle(box, radius=28, fill=fill)
        draw.text((x + 24, y + 24), item["title"], font=get_font(32, bold=True), fill=(30, 43, 61))
        wrapped = wrap_text(draw, item["goal"], body_font, 410)
        yy = y + 82
        for line in wrapped[:3]:
            draw.text((x + 24, yy), line, font=body_font, fill=(63, 77, 98))
            yy += 34
        if x > 600:
            x = 70
            y += 220
        else:
            x = 570

    draw.text((70, 820), "这张图作为该方向课程组图，可直接放在每份文档前半部分增强图文感。", font=sub_font, fill=(88, 99, 116))
    image.save(path)
    return path


def explain_bullet(topic: str, bullet: str) -> list[str]:
    return [
        f"{bullet}。放到真实学习场景里，这一条的重点不是把概念背下来，而是知道你在实际操作时应该把注意力放在哪里。很多人之所以越学越乱，就是因为知道名词，却不知道它在任务里承担什么角色。",
        f"围绕“{topic}”学习时，建议你不要只停留在理解表面意思，而是马上把它和自己的工作场景对应起来：它能帮你省哪一步时间，能降低哪一类重复劳动，哪里又必须继续由你自己判断。只有这样，这个知识点才会从‘我听过’变成‘我会用’。",
    ]


def add_detail_block(doc: Document, title: str, bullets: list[str], topic: str) -> None:
    doc.add_heading(title, level=2)
    for bullet in bullets:
        p = doc.add_paragraph()
        run = p.add_run(bullet)
        run.bold = True
        for line in explain_bullet(topic, bullet):
            add_paragraph(doc, line)


def add_case_block(doc: Document, topic: str) -> None:
    doc.add_heading("一个具体场景", level=1)
    add_paragraph(
        doc,
        f"假设你现在就要把“{topic}”用到一个真实任务里，最好的做法不是继续搜更多教程，而是先找一件最近 7 天内一定会重复出现的事。这个任务不需要很大，甚至可以只是写一条内容、整理一段资料、设计一个资料包标题。关键是它足够真实，足够近期，做完以后你能立刻感受到差别。",
    )
    add_paragraph(
        doc,
        "把真实任务拿出来以后，你会更容易判断：哪些部分适合让 AI 或 workflow 帮你提速，哪些部分必须由你自己拍板。真正有效的课程不是让你记住更多理论，而是让你在一个真实任务里第一次跑出手感。",
    )


def add_case_study_block(doc: Document, topic: str) -> None:
    doc.add_heading("案例拆解", level=1)
    add_paragraph(
        doc,
        f"下面用一个更完整的案例来理解“{topic}”。假设你是一个准备做内容和私域承接的普通创作者，你不是技术人员，也没有大团队，但你希望用更少时间做出更稳定的内容输出。这个时候，你最不需要的就是继续囤更多工具；你更需要的是先抓住一条清晰路径。",
    )
    add_paragraph(
        doc,
        "案例里的关键动作通常是这样的：先把目标说清楚，再把任务拆成可执行步骤，然后只挑最重复、最耗时、最标准化的那几步交给 AI 或 workflow 帮你提速。最后再由自己做判断、修正和发布。很多人觉得自己学了很久还是没效果，真正原因往往不是不会，而是从来没按这个顺序认真跑过一次。",
    )
    add_paragraph(
        doc,
        "如果你把这个案例代入自己的现实，你会发现很多知识点一下就顺了：为什么要先理解概念，为什么不能只要 prompt 模板，为什么需要 workflow，为什么 demo 看起来很厉害但结果还得靠训练和反馈。这也是课程真正想让你明白的地方。",
    )


def add_faq_block(doc: Document, topic: str) -> None:
    doc.add_heading("客户常见提问", level=1)
    faq_pairs = [
        (
            f"我是不是把“{topic}”学懂了，就能马上看到结果？",
            "不一定。学懂这一节，通常意味着你开始知道应该怎么做，而不是已经稳定做到。真正的变化来自你把它放进真实任务里，连续做几次，再根据结果调整。对大多数人来说，最大的进步不是多学会一个概念，而是第一次把正确动作跑顺。",
        ),
        (
            "我是不是还要继续学很多工具？",
            "多数情况下不用。先把手上最常见的一两个工具用顺，比继续囤十个新工具更有价值。真正会拉开差距的是任务拆解、输入质量和复盘能力，而不是工具数量。",
        ),
        (
            "为什么我看懂了，真正做的时候还是卡住？",
            "因为理解和执行中间还差一层练习。很多人以为自己缺的是新资料，其实缺的是把同一个动作连续做几次，并有人帮自己指出哪里做偏了。也正因为这样，很多知识资料适合入门，而更强反馈的学习形式更适合拿结果。",
        ),
    ]
    for q, a in faq_pairs:
        add_paragraph(doc, q, bold_prefix=q)
        add_paragraph(doc, a)


def add_analogy_block(doc: Document, topic: str) -> None:
    doc.add_heading("用生活中的例子理解", level=1)
    analogies = [
        f"如果把“{topic}”放到做饭里理解：你不会因为买了一个更贵的锅，就自动做出一桌稳定好吃的菜。真正决定结果的，是菜谱顺序、食材准备、火候判断和你做过多少次。AI、agent、skill 这些东西更像厨房里的不同工具；workflow 更像菜谱顺序；而训练和反馈更像你一次次试菜、调味和改进。",
        "再换一个生活场景：学开车的时候，你不可能只看说明书就上路很稳。你先要知道方向盘、刹车、后视镜分别干什么，但真正变熟，是靠一遍遍倒车、转弯、看路况。AI 学习也是一样，概念当然要懂，但真正有效的是在真实路况里不断修正动作。",
        "所以很多人觉得自己学了很多却没成果，本质上不是信息不够，而是一直停在“看说明书”的阶段，没有进入“真正把动作练顺”的阶段。",
    ]
    for text in analogies:
        add_paragraph(doc, text)


def add_mistakes_block(doc: Document, topic: str) -> None:
    doc.add_heading("常见误区", level=1)
    add_bullets(
        doc,
        [
            f"学“{topic}”时一上来就追求一步到位，结果把任务做得过大。",
            "以为看懂内容就等于会操作，没有把知识点落到自己的实际任务里。",
            "只收藏模板和工具，不记录自己的输入、修改和结果，所以越学越散。",
        ],
    )
    add_paragraph(
        doc,
        "如果你想避免这些问题，一个很简单的标准是：每学完一节，都要留下一个你自己真的会继续使用的动作、模板或流程片段。没有留下任何可复用成果，说明这节内容还没有真正进入你的工作系统。",
    )


def add_steps_block(doc: Document, topic: str, actions: list[str]) -> None:
    doc.add_heading("建议你这样开始", level=1)
    add_numbered(doc, actions)
    add_paragraph(
        doc,
        f"这组动作的设计原则很简单：先让你把“{topic}”跑起来，再考虑是否做复杂升级。很多人失败不是因为不会高级技巧，而是因为从来没有认真跑通过第一条最短路径。",
    )


def add_checklist_block(doc: Document, topic: str) -> None:
    doc.add_heading("学完这节后的自查清单", level=1)
    add_bullets(
        doc,
        [
            f"我能不能用自己的话讲清楚“{topic}”是什么。",
            "我能不能说出它最适合解决哪类任务，而不是泛泛而谈。",
            "我有没有把今天的内容变成一个可继续复用的动作、模板或流程。",
            "我知不知道下一步应该继续练什么，而不是又回去到处搜工具。",
        ],
    )


def add_upgrade_block(doc: Document) -> None:
    doc.add_heading("如果你想继续学下去", level=1)
    add_paragraph(
        doc,
        "到这里为止，这份文档的任务是帮你先建立判断、先跑出最小动作。再往下走，真正决定结果的，不是再多看几份资料，而是有没有人帮你看输入、看流程、看输出，帮你把错误改掉，把动作练熟。所以如果你已经知道自己卡在哪一步，后续最有效的学习方式通常不再是继续囤资料，而是进入更强反馈的学习环境。",
    )


def write_doc(category: dict, item: dict) -> Path:
    doc = Document()
    configure_document(doc)
    add_title_block(doc, item["title"], item["subtitle"], item["goal"])
    add_paragraph(
        doc,
        "你拿到这份内容以后，不需要先追求全部理解，更不要急着继续搜更多资料。最稳的方式，是先看懂这一节想帮你解决什么问题，再把里面的一两个动作放到自己的真实任务里。只要你真的动手，这份文档的价值就会立刻显现出来。",
    )

    doc.add_heading("适合谁", level=1)
    add_bullets(doc, item["for_whom"])

    doc.add_heading("学完你会得到什么", level=1)
    add_paragraph(
        doc,
        item["goal"] + " 这不是一份为了堆知识点而写的材料，而是一份希望你看完之后就知道下一步该怎么做的学习文档。",
    )

    img = image_path(item["image"])
    category_banner = CATEGORY_IMG_DIR / f"{category['folder']}.png"
    if category_banner.exists():
        add_image(doc, category_banner, "本方向课程导图")
    if img and img.exists():
        add_image(doc, img, "辅助讲解图")

    doc.add_heading("为什么这一节值得先学", level=1)
    add_paragraph(
        doc,
        f"很多人学习“{item['title']}”时，会把注意力放在名词、热点和工具上，但真正拉开差距的往往不是知道得更多，而是能不能更快把概念转成动作。你只要比别人更早把它放进自己的任务流程里，就已经领先了一步。",
    )
    add_paragraph(
        doc,
        "对普通用户来说，学习顺序比学习数量更重要。先学对，再学多，效果完全不一样。这节内容存在的意义，就是帮你把顺序排对。",
    )

    doc.add_heading("核心内容", level=1)
    for section_title, bullets in item["sections"]:
        add_detail_block(doc, section_title, bullets, item["title"])

    add_case_block(doc, item["title"])
    add_case_study_block(doc, item["title"])
    add_mistakes_block(doc, item["title"])
    add_faq_block(doc, item["title"])
    add_analogy_block(doc, item["title"])
    add_steps_block(doc, item["title"], item["actions"])
    add_checklist_block(doc, item["title"])
    add_upgrade_block(doc)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("把这份内容真正用到你的任务里，效果会比你继续囤十份资料更明显。")
    footer_run.font.size = Pt(9)

    path = OUT_DIR / category["folder"] / item["name"]
    doc.save(path)
    return path


def write_category_index(category: dict, built_paths: list[Path]) -> None:
    doc = Document()
    configure_document(doc)
    add_title_block(doc, category["title"], category["intro"], "每个方向先给你 6 份独立资料，方便拆成课程、资料包、直播稿和私域答疑。")
    add_paragraph(doc, "使用建议：不要一次性全部发给用户。先按用户类型发其中 1 到 2 份，再根据反馈把人导向系统课或训练营。")
    doc.add_heading("本方向资料目录", level=1)
    rows = []
    for path in built_paths:
        rows.append([path.stem, path.name, "可单独发，也可组合成一个专题资料包"])
    build_table(doc, ["标题", "文件名", "用途"], rows, [2.0, 2.0, 2.5])
    doc.add_heading("最适合的承接动作", level=1)
    add_bullets(
        doc,
        [
            "把其中 1 份做成公开资料领取入口。",
            "把其中 2 到 3 份组合成低客单产品。",
            "把全部文档按顺序串起来，就能变成系统课素材底稿。",
        ],
    )
    doc.save(OUT_DIR / category["folder"] / "_本方向资料总目录.docx")


def write_root_index(all_paths: list[tuple[str, Path]]) -> None:
    doc = Document()
    configure_document(doc)
    add_title_block(doc, "AI 系统课程资料库总目录", "按 5 个方向拆分，每个方向 6 份独立 docx", "适合公开引流、私域承接、低客单产品和训练营课件准备")
    add_paragraph(doc, "这套资料库优先服务你现在的目标：让用户先看懂、先加你、先拿到 demo，再顺势进入系统课和训练营。")
    doc.add_heading("资料库总览", level=1)
    rows = []
    for folder_name, path in all_paths:
        rows.append([folder_name, path.stem, path.name])
    build_table(doc, ["方向", "文档标题", "文件名"], rows, [1.8, 2.3, 2.4])
    doc.add_heading("你可以怎么用", level=1)
    add_numbered(
        doc,
        [
            "把其中最基础的文档做成免费领取资料。",
            "把同一方向的 3 到 6 份资料打包成低客单专题。",
            "把 5 个方向串起来，直接升级成系统课程主干。",
            "把每份文档后面的‘学后动作’变成训练营作业。",
        ],
    )
    doc.save(OUT_DIR / "00_AI系统课程资料库总目录.docx")


def main() -> None:
    ensure_dirs()
    ensure_output()

    # Ensure shared images exist
    build_diagram_concept_map()
    build_diagram_content_flow()
    build_diagram_funnel()
    build_diagram_learning_path()
    for category in CATEGORIES:
        build_category_banner(category)

    all_paths: list[tuple[str, Path]] = []
    for category in CATEGORIES:
        built: list[Path] = []
        for item in category["docs"]:
            built_path = write_doc(category, item)
            built.append(built_path)
            all_paths.append((category["folder"], built_path))
        write_category_index(category, built)

    write_root_index(all_paths)

    readme = OUT_DIR / "README.txt"
    with open(readme, "w", encoding="utf-8") as fh:
        fh.write("AI 系统课程资料库 v2 已生成。\n\n")
        for category in CATEGORIES:
            fh.write(f"{category['folder']}\\\n")
            fh.write(f"  - _本方向资料总目录.docx\n")
            for item in category["docs"]:
                fh.write(f"  - {item['name']}\n")
            fh.write("\n")


if __name__ == "__main__":
    main()
