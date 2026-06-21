from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_course_materials import (
    ROOT,
    add_bullets,
    add_image,
    add_numbered,
    add_paragraph,
    add_title_block,
    build_diagram_concept_map,
    build_diagram_content_flow,
    build_diagram_funnel,
    build_diagram_learning_path,
    configure_document,
    get_font,
    wrap_text,
)


OUT_DIR = ROOT / "deliverables" / "course_library_v3"
GEN_DIR = OUT_DIR / "generated_images"
GEN_PNG_DIR = OUT_DIR / "generated_images_png"
LOCAL_DIR = OUT_DIR / "_local_images"
TOPIC_DIR = OUT_DIR / "_topic_images"


COURSE_GROUPS = [
    {
        "folder": "01_认知基础课",
        "title": "认知基础课",
        "subtitle": "先把 AI、大模型、Agent、Workflow、Skill 的关系讲明白",
        "cover_prefix": "foundations-cover",
        "docs": [
            ("01_AI到底是什么", "把 AI 从热点词变成你能讲给客户听清楚的工作概念", ["AI 不是什么魔法，它更像一个加速器。", "AI 最先该解决的不是炫技，而是重复劳动。", "客户真正关心的是它能不能帮自己省时间、提效率、扩选择。"]),
            ("02_大模型和普通工具的区别", "让用户知道大模型是底层能力，不是某个单一软件", ["大模型像发动机，应用像装了不同车身的车。", "别把产品名当能力本身。", "先理解能力边界，后面选工具就不会乱。"]),
            ("03_Agent为什么比聊天更进一步", "把 Agent 讲成会理解目标并调工具的执行者", ["聊天是回答，Agent 更偏执行。", "一旦涉及搜索、写文件、操作网页，Agent 的价值就会很明显。", "很多人学不会，是因为没有真实任务。"]),
            ("04_Workflow为什么比工具更重要", "先学流程，再学工具，才不会越学越散", ["Workflow 决定步骤顺序。", "工具只是装在步骤里的零件。", "先有任务，再决定上什么工具。"]),
            ("05_Skill怎么帮助你快速做Demo", "Skill 的本质是把经验和动作封装成可复用模块", ["Skill 适合快速打样。", "Skill 不等于结果保证。", "真正的稳定输出还是要靠训练与反馈。"]),
        ],
    },
    {
        "folder": "02_模型与提示实操课",
        "title": "模型与提示实操课",
        "subtitle": "把提问、提示词、角色设定、输出格式变成可操作能力",
        "cover_prefix": "prompting-cover",
        "docs": [
            ("01_新手提问法", "先学会把模糊需求讲清楚，再谈高级技巧", ["目标、对象、背景、输出格式，这四件事先说清楚。", "不会问，后面所有自动化都会不稳定。", "最好的提示词不是花哨，而是清楚。"]),
            ("02_结构化提示模板", "给你一套长期可复用的提示骨架", ["角色、任务、限制、格式，是最常用的四层结构。", "结构越清楚，返工越少。", "模板不是为了收藏，而是为了复用。"]),
            ("03_长文写作怎么和AI配合", "把 AI 从写稿器变成写作协作者", ["先由人定观点，再让 AI 补结构。", "不要一开始就让 AI 直接写终稿。", "分轮次输出，稳定性远高于一轮成稿。"]),
            ("04_信息整理与总结提示", "普通人最容易最快感受到 AI 价值的一类任务", ["长文摘要、评论归类、会议纪要都很适合 AI。", "越混乱的原始材料，越需要先给整理标准。", "先拿整理类任务建立信心，最稳。"]),
            ("05_角色提示与风格控制", "减少 AI 味，重点是上下文和标准而不是形容词", ["角色提示能稳定视角。", "样例和禁忌比华丽词藻更有用。", "AI 不会凭空拥有你的个人经历。"]),
        ],
    },
    {
        "folder": "03_Agent与自动化实操课",
        "title": "Agent与自动化实操课",
        "subtitle": "让用户从会问升级到会跑流程、会调用工具",
        "cover_prefix": "agent-automation-cover",
        "docs": [
            ("01_Agent最小结构", "目标、模型、工具、状态、输出，是最小可讲清楚的结构", ["别把 Agent 讲复杂。", "先让用户知道它为什么能分步执行。", "理解结构后，案例才不会只是热闹。"]),
            ("02_工具调用和函数调用", "让用户知道 Agent 为什么能真正做事", ["函数调用是让模型给出结构化参数。", "工具层是 Agent 的手和脚。", "没有工具层，就很难从回答问题走向完成任务。"]),
            ("03_浏览器自动化能做什么", "把网页操作类自动化讲成半自动执行助手", ["适合做登录后页面操作和草稿提交。", "最适合演示，而不是夸大全自动生意。", "关键动作仍然建议保留人工确认。"]),
            ("04_怎么把任务拆成自动化流程", "一切自动化都从任务拆解开始", ["先拆步骤，再想自动化。", "最适合交给 AI 的是重复、标准、耗时的环节。", "不要在没有流程图的时候谈自动化。"]),
            ("05_公开Agent案例怎么学", "学案例不是为了照搬，而是为了打开思路", ["内容 agent、研究 agent、浏览器 agent 都值得看。", "看案例要学结构，不是只看效果图。", "真正适合你的，永远要落回自己的任务。"]),
        ],
    },
    {
        "folder": "04_内容创作实操课",
        "title": "内容创作实操课",
        "subtitle": "让课程从认知真正落到选题、写作、改写、发布、复盘",
        "cover_prefix": "content-practical-cover",
        "docs": [
            ("01_选题Workflow", "先拆人群和问题，再决定写什么", ["选题不是灵感问题，而是人群和场景问题。", "AI 最适合帮你拆痛点和列角度。", "没有承接动作的选题，很难转化。"]),
            ("02_标题和开场钩子", "让用户停下来，不靠堆砌夸张词", ["标题负责筛点击。", "开场负责留住人。", "同一主题要会写多个角度。"]),
            ("03_提纲到初稿的协作", "先由人定骨架，再让 AI 补表达", ["观点要自己定。", "AI 更适合扩写、补句子、补案例。", "骨架对了，后面返工会少很多。"]),
            ("04_多平台改写", "同一主题改成公众号、小红书、朋友圈不同版本", ["不是复制，是重写。", "保留观点，重构表达。", "AI 最适合做不同平台的语气和长度切换。"]),
            ("05_发布前检查与复盘", "提效不只是生成，还包括检查和复盘", ["发布前要校对 CTA、错别字和口径。", "发布后要看私信、收藏、评论和加你比例。", "没有复盘，后面所有流程都不会进化。"]),
        ],
    },
    {
        "folder": "05_私域转化实操课",
        "title": "私域转化实操课",
        "subtitle": "把公开内容、免费资料、低客单、训练营串成闭环",
        "cover_prefix": "private-domain-cover",
        "docs": [
            ("01_免费资料包怎么设计", "资料包不是白送内容，而是筛选器和入口", ["最重要的是让用户愿意加你。", "资料包主题要和内容主题一一对应。", "越能引出下一步，资料包价值越高。"]),
            ("02_留资与加微信路径", "内容不是终点，进入私域才是真正开始", ["CTA 要单一。", "加你动作要低门槛。", "不要一条内容放三个不同目标。"]),
            ("03_低客单产品怎么做", "低客单的作用不是赚钱最多，而是验证需求和筛人", ["模板包、微课、清单包都很适合。", "低客单是升级系统课和训练营的桥。", "卖的是确定的下一步，而不是堆内容。"]),
            ("04_系统课怎么搭", "把碎片知识整理成递进课程结构", ["先认知，再工具，再流程，再案例，再转化。", "不要一上来做大而全。", "从真实问题池反推课纲，最稳。"]),
            ("05_训练营为什么一定要有", "训练营的价值在于纠偏、作业、反馈和结果", ["资料负责打开认知。", "系统课负责建立方法。", "训练营负责把动作练顺。"]),
        ],
    },
    {
        "folder": "06_赠送Skills与Workflow包",
        "title": "赠送Skills与Workflow包",
        "subtitle": "让用户学完不只是听懂，还能立刻拿到可用模板和流程包",
        "cover_prefix": "gift-pack-cover",
        "docs": [
            ("01_内容选题Skill包", "直接可用的选题拆解、角度扩展、评论归类思路包", ["适合拿来快速做内容前期准备。", "重点是先有输入，再跑模板。", "不要把模板当答案，要把它当脚手架。"]),
            ("02_长文写作Workflow包", "从题目到提纲到扩写到改写的最小工作流", ["适合公众号、课程文稿、长图文。", "先由人定骨架。", "再让 AI 分步补全。"]),
            ("03_多平台改写Workflow包", "同一主题改成多个平台版本的流程包", ["保留主观点。", "改结构和语气。", "改 CTA 而不是照搬。"]),
            ("04_私域资料包设计清单", "一份能直接拿来设计免费资料和低客单产品的清单", ["适合引流和承接。", "帮助你避免大而全。", "先解决一个小结果。"]),
            ("05_训练营作业模板包", "把训练营从讲知识变成带动作的模板集合", ["包含作业、检查、复盘三个层次。", "适合陪跑营和实战营。", "让用户每学一节都能交出东西。"]),
        ],
    },
    {
        "folder": "07_7天实战训练营版",
        "title": "7天实战训练营版",
        "subtitle": "把前面的认知、实操、转化全部压缩成 7 天可执行动作",
        "cover_prefix": "bootcamp-cover",
        "docs": [
            ("01_第1天_定位与目标", "第一天先不追求产出，先把人群、目标和学习结果钉住", ["先明确你服务谁。", "先确定 7 天后要拿到什么结果。", "目标越具体，后面动作越不飘。"]),
            ("02_第2天_选题与资料包方向", "第二天把公开内容主题和免费资料包主题对齐", ["内容题目要能自然引出资料包。", "资料包要能自然引出私域。", "别做大而全，先解决一个小问题。"]),
            ("03_第3天_提示词与内容初稿", "第三天开始真正动手，用 AI 帮你拆题、写提纲、出初稿", ["先写提纲。", "再让 AI 扩写。", "最后自己改成能发布的版本。"]),
            ("04_第4天_多平台改写与发布", "第四天把同一主题变成多平台版本并安排发布", ["公众号、小红书、朋友圈不要照搬。", "保留观点，重写表达。", "每个平台只保留一个 CTA。"]),
            ("05_第5天_私域承接与低客单设计", "第五天把公开内容真正接到私域和低客单产品上", ["没有承接，流量就会散。", "低客单不是为了赚最多，而是为了筛人。", "先打通最小闭环。"]),
        ],
    },
    {
        "folder": "08_可直接复制使用包",
        "title": "可直接复制使用包",
        "subtitle": "不是讲概念，而是把你现在就能用的模板、清单、脚本整理出来",
        "cover_prefix": "template-pack-cover",
        "docs": [
            ("01_选题模板50条", "直接可改写的选题模板，用来快速起草内容方向", ["按照人群、痛点、场景、结果四个角度拆。", "不要只换词，要换切入点。", "模板是起点，不是终点。"]),
            ("02_标题模板与钩子模板", "直接可套的标题和开场模板，适合短内容与图文内容", ["标题负责停留。", "开场负责继续看。", "不要只会一种句式。"]),
            ("03_长文提纲与扩写模板", "适合公众号、课程稿、长图文的结构化写作模板", ["先写骨架。", "再补案例。", "最后统一语气。"]),
            ("04_私域承接文案模板", "用于评论区、私信、朋友圈、公众号回复的承接文案", ["动作要单一。", "语气要自然。", "先帮用户迈出下一步。"]),
            ("05_训练营作业与复盘清单", "把训练营真正做出结果的作业模板和复盘模板整理好", ["每一节课都要交付一个动作。", "每一次作业都要有复盘。", "让用户看见自己的进步。"]),
        ],
    },
    {
        "folder": "09_模板填空版",
        "title": "模板填空版",
        "subtitle": "把课程内容真正变成客户可以直接复制、直接改、直接发的成品模板",
        "cover_prefix": "template-fill-cover",
        "docs": [
            ("01_账号定位填空模板", "一页一页填空，快速收敛定位、人群、承接动作", ["先确定你服务谁。", "再确定你解决什么问题。", "最后确定内容和承接怎么接起来。"]),
            ("02_免费资料包填空模板", "把资料包题目、大纲、CTA、承接路径做成可填写表", ["让用户一边填，一边把产品想清楚。", "避免做成大而全资料。", "先解决一个小结果。"]),
            ("03_内容选题与标题填空模板", "让用户围绕人群、痛点、结果写出一批可用选题", ["不是灵感表，而是结构化选题表。", "标题不是单独写，要和场景一起写。", "每个标题都要知道要承接什么动作。"]),
            ("04_私域承接文案填空模板", "把评论区、私信、朋友圈、公众号回复做成现成骨架", ["少写废话。", "先帮用户迈出下一步。", "每条文案只保留一个动作目标。"]),
            ("05_训练营作业布置填空模板", "把每节课的作业目标、提交格式、复盘问题做成统一模板", ["让训练营交付更整齐。", "让学员更容易动起来。", "让你更容易批改和复盘。"]),
        ],
    },
    {
        "folder": "10_训练营作业与打卡版",
        "title": "训练营作业与打卡版",
        "subtitle": "直接拿去用的作业页、打卡页、复盘页，帮助训练营真正跑起来",
        "cover_prefix": "bootcamp-worksheet-cover",
        "docs": [
            ("01_开营说明与目标设定", "让学员在开营当天就知道自己接下来 7 天要完成什么", ["先定结果。", "再定动作。", "最后定提交方式。"]),
            ("02_第1到2天作业页", "围绕定位、选题、资料包方向做第一次输出", ["先出草稿。", "不求完美，先交作业。", "用提交动作逼出真实进度。"]),
            ("03_第3到4天作业页", "围绕提示词、初稿、改写和发布开始真正动手", ["开始让 AI 进入流程。", "开始形成第一批内容资产。", "开始练从想法到成稿。"]),
            ("04_第5到7天作业页", "围绕承接、低客单、训练营逻辑完成闭环动作", ["不只是发内容。", "还要让内容接上私域。", "还要让私域接上产品。"]),
            ("05_每日打卡与周复盘页", "让训练营不只是上课，而是持续推进和复盘", ["每天有动作。", "每天有提交。", "每周有复盘。"]),
        ],
    },
    {
        "folder": "11_成交与预热文案包",
        "title": "成交与预热文案包",
        "subtitle": "把课程卖出去前最需要的销售页、朋友圈、小红书和私域文案补齐",
        "cover_prefix": "sales-copy-cover",
        "docs": [
            ("01_销售页长文案", "一份可直接改写成详情页或长图文的成交文案框架", ["先讲问题。", "再讲路径。", "再讲你为什么能带他做到。"]),
            ("02_朋友圈预热文案包", "适合连续 5 到 7 天预热的朋友圈文案结构", ["先打认知。", "再打需求。", "最后才打产品。"]),
            ("03_小红书引流文案包", "适合拿来改成图文贴或评论区承接文案", ["标题要强。", "正文要轻。", "CTA 要单一。"]),
            ("04_私域答疑成交文案", "常见犹豫点、顾虑点、比较点的标准答法", ["不要空泛安慰。", "要把结果路径讲清楚。", "要让客户知道下一步很明确。"]),
            ("05_开营通知与转介绍文案", "用于开营、提醒、续费、转介绍的关键节点文案", ["别只会卖一次。", "要让课程结束后还能继续转化。", "每次通知都带一个明确动作。"]),
        ],
    },
]

COURSE_GROUPS.extend(
    [
        {
            "folder": "12_SOP标准操作课",
            "title": "SOP标准操作课",
            "subtitle": "把常见任务拆成标准动作，不再只停留在知道概念",
            "cover_prefix": "sop-practical-cover",
            "docs": [
                ("01_选题SOP怎么写", "把选题从凭感觉改成按步骤推进的标准动作", ["先写服务对象，再写问题场景，再写结果承接。", "先做出一个最小版本，再扩成批量版本。", "SOP 的重点不是好看，而是别人照着也能走。"]),
                ("02_长文写作SOP", "把一篇长文从想法到成稿拆成可执行步骤", ["先定主判断。", "再列提纲。", "最后才让 AI 进入扩写和润色。"]),
                ("03_私域资料包SOP", "把免费资料从想法变成真正能承接加你的资料包", ["先定关键词。", "再定目录。", "再定加你以后的跟进动作。"]),
                ("04_训练营开营SOP", "把开营准备、通知、作业、答疑做成可重复执行动作", ["先准备交付节奏。", "再准备打卡机制。", "最后再讲开营氛围和转化。"]),
                ("05_复盘优化SOP", "把复盘从泛泛总结变成下一轮可直接改的动作", ["只看关键指标。", "只找最卡一步。", "只改最影响结果的一个点。"]),
            ],
        },
        {
            "folder": "13_Workflow设计课",
            "title": "Workflow设计课",
            "subtitle": "把任务串成流程，让 AI、模板、人工判断各归各位",
            "cover_prefix": "workflow-design-cover",
            "docs": [
                ("01_从任务到Workflow", "先把真实任务拆开，再决定怎么串成流程", ["任务没拆开，流程一定空。", "流程图不是形式，而是防止返工。", "先跑通最短流程，再考虑加工具。"]),
                ("02_内容生产Workflow", "把选题、提纲、初稿、改写、发布串成一条线", ["每一步只产出一个明确结果。", "AI 负责提速，不负责替你做判断。", "流程越清楚，内容越稳定。"]),
                ("03_私域承接Workflow", "把公开内容、私信、资料包、低客单接成闭环", ["先定入口。", "再定承接。", "最后定成交后的下一步。"]),
                ("04_训练营交付Workflow", "把开营、作业、答疑、复盘做成连续动作", ["别只安排内容。", "更要安排提交和反馈。", "流程要让学员一直有下一步。"]),
                ("05_多人协作Workflow", "把一个人能跑通的动作扩成团队可协作流程", ["先统一输入标准。", "再统一输出格式。", "最后统一复盘口径。"]),
            ],
        },
        {
            "folder": "14_MCP与工具接入课",
            "title": "MCP与工具接入课",
            "subtitle": "把工具接入、能力调用、半自动执行讲清楚，不再只是听概念",
            "cover_prefix": "mcp-tools-cover",
            "docs": [
                ("01_MCP到底是什么", "把 MCP 讲成可接能力的标准接口，而不是玄乎的新名词", ["先理解它为什么存在。", "再理解它解决的是接入问题。", "最后才谈怎么为业务所用。"]),
                ("02_怎么接浏览器与文档能力", "理解浏览器、文档、表格这类能力是怎么接进流程的", ["先确认任务。", "再确认工具能力。", "最后确认哪些动作必须人工确认。"]),
                ("03_MCP配置案例", "给出能直接改的配置片段，不再停留在原理层", ["先跑最小可用配置。", "再加业务字段。", "最后再整理成团队标准模板。"]),
                ("04_MCP常见报错排查", "把常见接入问题拆成一套排查动作", ["先看路径。", "再看权限。", "再看参数。"]),
                ("05_MCP能接进哪些业务", "把工具接入落回选题、写作、资料交付、训练营这些真实任务", ["不是为了接而接。", "而是为了让任务更顺。", "能省时间的地方才值得接。"]),
            ],
        },
        {
            "folder": "15_交付与排错实战课",
            "title": "交付与排错实战课",
            "subtitle": "把最后一公里补齐，教会怎么改、怎么查、怎么交付出去",
            "cover_prefix": "delivery-debug-cover",
            "docs": [
                ("01_为什么做了很多还是交付不出来", "把卡在最后一步的常见问题讲透", ["问题常常不在不会做，而在没有收口。", "没有交付格式，动作就容易散。", "先做出可交版本，再谈完美。"]),
                ("02_怎么改成自己的版本", "把模板、文案、流程改成自己的业务表达", ["先改对象。", "再改场景。", "最后改 CTA 和结果。"]),
                ("03_常见故障排查方法", "把内容、流程、配置、工具这些故障统一成排查思路", ["先缩小范围。", "再排除变量。", "最后记录结论。"]),
                ("04_怎么给客户交付", "把资料、课程、模板、作业页变成真正可发出去的成品", ["交付要有命名。", "交付要有说明。", "交付要有下一步动作。"]),
                ("05_怎么把一套资料持续迭代", "让课程资料不是一次性产物，而是越来越好用的资产", ["每轮只改高频问题。", "每轮都补真实案例。", "每轮都留下更稳定版本。"]),
            ],
        },
    ]
)


LOCAL_IMAGE_MAP = {
    "concept": "01_ai_concept_map.png",
    "content": "02_content_workflow.png",
    "funnel": "03_private_domain_funnel.png",
    "learning": "04_learning_paths.png",
}


GROUP_FALLBACK_IMAGE = {
    "01_认知基础课": "concept",
    "02_模型与提示实操课": "concept",
    "03_Agent与自动化实操课": "learning",
    "04_内容创作实操课": "content",
    "05_私域转化实操课": "funnel",
    "06_赠送Skills与Workflow包": "learning",
    "07_7天实战训练营版": "funnel",
    "08_可直接复制使用包": "content",
    "09_模板填空版": "learning",
    "10_训练营作业与打卡版": "funnel",
    "11_成交与预热文案包": "content",
    "12_SOP标准操作课": "content",
    "13_Workflow设计课": "learning",
    "14_MCP与工具接入课": "concept",
    "15_交付与排错实战课": "funnel",
}


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    GEN_DIR.mkdir(parents=True, exist_ok=True)
    GEN_PNG_DIR.mkdir(parents=True, exist_ok=True)
    LOCAL_DIR.mkdir(parents=True, exist_ok=True)
    TOPIC_DIR.mkdir(parents=True, exist_ok=True)
    for group in COURSE_GROUPS:
        (OUT_DIR / group["folder"]).mkdir(parents=True, exist_ok=True)


def generate_local_banner(group: dict) -> Path:
    path = LOCAL_DIR / f"{group['folder']}.png"
    image = Image.new("RGB", (1600, 900), (248, 246, 241))
    draw = ImageDraw.Draw(image)
    title_font = get_font(54, bold=True)
    sub_font = get_font(28)
    body_font = get_font(30)
    draw.text((70, 60), group["title"], font=title_font, fill=(34, 44, 61))
    for i, line in enumerate(wrap_text(draw, group["subtitle"], sub_font, 1380)):
        draw.text((70, 145 + i * 42), line, font=sub_font, fill=(98, 109, 126))

    colors = [(255, 233, 214), (252, 241, 205), (220, 233, 250), (220, 239, 228), (232, 225, 248)]
    y = 280
    x = 70
    for idx, doc_item in enumerate(group["docs"]):
        color = colors[idx % len(colors)]
        draw.rounded_rectangle((x, y, x + 280, y + 220), radius=24, fill=color)
        draw.text((x + 22, y + 20), doc_item[0].split("_", 1)[1], font=get_font(24, bold=True), fill=(31, 43, 61))
        lines = wrap_text(draw, doc_item[1], body_font, 240)
        yy = y + 78
        for line in lines[:4]:
            draw.text((x + 22, yy), line, font=body_font, fill=(63, 77, 98))
            yy += 32
        x += 300
        if idx == 1:
            x = 70
            y += 250
        elif idx == 3:
            x = 370
    image.save(path)
    return path


def generate_topic_banner(group: dict, stem: str, title: str, promise: str, bullets: list[str]) -> Path:
    path = TOPIC_DIR / f"{group['folder']}_{stem}.png"
    image = Image.new("RGB", (1600, 900), (246, 244, 239))
    draw = ImageDraw.Draw(image)
    title_font = get_font(52, bold=True)
    sub_font = get_font(28)
    body_font = get_font(26)
    draw.text((70, 60), title, font=title_font, fill=(31, 43, 61))
    for idx, line in enumerate(wrap_text(draw, promise, sub_font, 1380)):
        draw.text((70, 145 + idx * 40), line, font=sub_font, fill=(88, 99, 116))

    palette = [(255, 232, 214), (252, 241, 206), (221, 234, 251)]
    positions = [(80, 300), (560, 300), (1040, 300)]
    for idx, bullet in enumerate(bullets[:3]):
        x, y = positions[idx]
        draw.rounded_rectangle((x, y, x + 420, y + 300), radius=28, fill=palette[idx % len(palette)])
        draw.text((x + 24, y + 24), f"重点 {idx + 1}", font=get_font(24, bold=True), fill=(30, 43, 61))
        lines = wrap_text(draw, bullet, body_font, 370)
        yy = y + 84
        for line in lines[:6]:
            draw.text((x + 24, yy), line, font=body_font, fill=(63, 77, 98))
            yy += 34

    footer_text = f"{group['title']} | {title}"
    draw.text((70, 790), footer_text, font=sub_font, fill=(88, 99, 116))
    image.save(path)
    return path


def get_generated_cover(prefix: str) -> Path | None:
    matches = sorted(GEN_DIR.glob(f"{prefix}-*.jpeg"))
    if matches:
        src = matches[-1]
        target = GEN_PNG_DIR / f"{src.stem}.png"
        if not target.exists() or target.stat().st_mtime < src.stat().st_mtime:
            Image.open(src).convert("RGB").save(target, format="PNG")
        return target
    return None


def build_story_sections(doc: Document, title: str, promise: str, bullets: list[str]) -> None:
    doc.add_heading("这节课先帮你解决什么问题", level=1)
    add_paragraph(doc, promise + " 你不需要一口气把它全部学完，更重要的是知道这节内容解决的是哪一类真实卡点。")
    add_paragraph(doc, "大多数人学不下去，不是因为内容太难，而是因为不知道这节课和自己当前任务有什么关系。所以这份讲义会一直围绕真实场景来写，而不是只堆概念。")

    doc.add_heading("核心讲解", level=1)
    for bullet in bullets:
        p = doc.add_paragraph()
        r = p.add_run(bullet)
        r.bold = True
        add_paragraph(doc, f"{bullet}。如果你只是把这句话看成一个知识点，它很快会忘；但如果你把它放进自己的工作里，你就会知道这一步到底帮你省了什么、挡住了什么坑。")
        add_paragraph(doc, "真正值钱的不是这句话本身，而是你学会用这句话来做判断。很多人看了很多资料，却始终做不出稳定结果，本质上就是没有把知识点转成判断标准。")


def add_case_study(doc: Document, title: str) -> None:
    doc.add_heading("案例拆解", level=1)
    add_paragraph(doc, f"假设你现在要把“{title}”用到一个真实项目里。你不是技术团队，也没有太多时间，但你需要更快地产出内容、设计资料、承接私域，或者跑一个最小的产品闭环。这个时候，最好的做法不是继续囤新工具，而是先把当下最常见的一件事跑顺。")
    add_paragraph(doc, "这个案例的重点在于：先缩小任务，再跑最短路径。比如先只解决一条内容、一份资料、一套简单模板，而不是一上来就想做全自动系统。很多人之所以学了很多还是没结果，就是因为一直在做超出自己当下阶段的大任务。")
    add_paragraph(doc, "当你把一个案例真正跑下来以后，后面的很多知识点自然都会串起来。你会知道哪里需要 AI，哪里需要自己判断，哪里需要模板，哪里需要训练。这种‘跑通一次带来的理解’，远比继续读三份新资料更有价值。")


def add_common_questions(doc: Document, title: str) -> None:
    doc.add_heading("客户常见提问", level=1)
    qas = [
        (
            f"我只学“{title}”这一节，够不够？",
            "这取决于你现在卡在哪一步。如果你只是需要先打开认知、跑通最小动作，这一节已经足够帮你迈出第一步。但如果你要的是稳定结果，后面一定还需要把动作练熟，并接受反馈和纠偏。",
        ),
        (
            "我是不是还要配很多工具一起学？",
            "大多数情况下不用。先把当前最关键的一两个动作跑顺，比同时装十个工具更重要。真正拉开差距的不是工具数量，而是你有没有一条能重复使用的流程。",
        ),
        (
            "为什么我看懂以后，做起来还是会卡？",
            "因为理解和执行是两回事。真正的跨越发生在你动手做、做完后复盘、复盘后再改。很多人以为自己缺的是更多知识，实际上缺的是连续练同一个动作。"),
    ]
    for q, a in qas:
        add_paragraph(doc, q, bold_prefix=q)
        add_paragraph(doc, a)


def add_plain_examples(doc: Document, title: str) -> None:
    doc.add_heading("用生活中的通俗例子理解", level=1)
    examples = [
        f"如果把“{title}”放到做饭里理解：你不会因为买了更贵的锅，就自动做出一桌稳定好吃的菜。真正决定结果的，是做菜顺序、备菜标准、火候判断，以及你做过多少次。AI 和各种工具更像锅、刀、灶台；流程更像菜谱顺序；训练和反馈更像你一次次试菜和调味。",
        "再换成开车的例子也一样：你知道方向盘、刹车和油门分别干什么，不代表你已经能在真实路况里开得很稳。真正变熟，是靠不断起步、转弯、看路况、修正动作。AI 学习也是同样的逻辑，理解工具功能不等于能稳定跑结果。",
        "所以当你觉得课程内容看起来都懂，但自己做不出来时，不要先怀疑是不是还缺一个神工具。更大概率是，你还停留在看说明书的阶段，没有进入把动作练顺的阶段。",
    ]
    for item in examples:
        add_paragraph(doc, item)


def add_practical_steps(doc: Document, title: str) -> None:
    doc.add_heading("课后实操动作", level=1)
    add_numbered(
        doc,
        [
            "找一个你最近 7 天内一定会重复出现的真实任务，不要选太大。",
            f"用“{title}”这节课的逻辑，把这个任务拆成 4 到 6 步。",
            "标出其中哪些步骤适合让 AI 或 workflow 介入，哪些必须自己判断。",
            "先跑一遍最短版本，不求完美，只求跑通。",
            "记下哪里最卡，下一轮只改一个点，不要一次改全部。",
        ],
    )
    add_paragraph(doc, "这套动作的核心不是做得多，而是做得真。只要你真的在一个真实任务里练过一次，这节课就已经开始产生价值。")


def add_takeaway_checklist(doc: Document, title: str) -> None:
    doc.add_heading("学完这节课后的检查清单", level=1)
    add_bullets(
        doc,
        [
            f"我能不能用自己的话讲清楚“{title}”到底解决什么问题。",
            "我能不能说出这节内容最适合放进哪类真实任务里。",
            "我有没有基于这节课留下一个自己会继续复用的模板、步骤或检查清单。",
            "我知不知道下一步该继续练哪个动作，而不是继续乱搜新工具。",
        ],
    )


def add_real_use_case(doc: Document, folder: str, title: str) -> None:
    doc.add_heading("实际案例：我会这样直接用", level=1)
    if folder == "04_内容创作实操课":
        add_paragraph(doc, "比如我现在要写一篇关于“为什么学了很多 AI 还是没提效”的图文。我不会一上来就让模型直接写成文，而是先自己写一句主判断：问题不在工具，而在流程。")
        add_paragraph(doc, "然后我会让模型先帮我拆出 5 个最容易共鸣的卡点，再把这 5 个卡点排成一个阅读顺序，最后给每个卡点补两种更口语的说法。我自己保留最后拍板的权力。")
        add_paragraph(doc, "这样做的结果，不是让我少思考，而是让我把思考放在更值钱的地方：决定观点、筛掉废话、保留真正能转化的表达。")
    elif folder == "05_私域转化实操课":
        add_paragraph(doc, "比如我要发一份免费资料包，我会先确定：这份资料包到底是给谁的、准备解决什么问题、看完以后我最想让对方做什么。")
        add_paragraph(doc, "如果我的目标是让对方加我，那我只保留一个动作，比如‘想拿完整资料，私信我关键词资料’。我不会同时让他评论、关注、加群、去公众号，因为那样只会把动作打散。")
        add_paragraph(doc, "对我来说，真正有效的承接从来都不是更花哨，而是更单一、更顺、更容易让对方迈出下一步。")
    elif folder in {"06_赠送Skills与Workflow包", "08_可直接复制使用包", "09_模板填空版"}:
        add_paragraph(doc, "如果我要把这一节做成可以送出去的模板或 workflow，我会先选一件自己已经真实做过很多次的事，而不是现想一个看起来很酷的例子。")
        add_paragraph(doc, "我先把自己真实做过的步骤写出来，再把这些步骤压缩成别人也能直接照着走的模板。这样交出去的东西才不会空，也不会只有表面结构。")
    else:
        add_paragraph(doc, "我实际使用这一节时，会先抓一个最近就要发生的小任务。只要这件事已经在我手上，我就能立刻判断这一节内容到底帮不帮得上忙。")
        add_paragraph(doc, "我不会先追求一次把所有环节都自动化，而是先让最耗时、最重复、最标准化的一步先顺起来。")


def add_how_to_modify(doc: Document, folder: str, title: str) -> None:
    doc.add_heading("拿到以后我会怎么改", level=1)
    add_paragraph(doc, "我不会直接整段照搬。最先改的一定是对象、场景和最后结果。对象不对，整个内容都会跑偏；场景不对，案例就会变空；结果不对，CTA 也会失效。")
    if folder in {"09_模板填空版", "11_成交与预热文案包"}:
        add_paragraph(doc, "如果这是模板或文案，我会先把里面所有泛词换成我自己的业务表达。比如把‘用户’改成‘内容创作者’，把‘转化’改成‘加我拿资料’，把‘产品’改成‘7天训练营’。")
        add_paragraph(doc, "换词以后，我会再删掉不属于我业务的句子。模板真正好用的标准，不是写得多完整，而是改完以后像我自己。")
    else:
        add_paragraph(doc, "如果这是方法课，我会先保留结构，不急着保留原句。结构是骨架，原句只是示范。只要骨架还在，我就能换成更适合自己的表达。")


def add_problem_solving(doc: Document) -> None:
    doc.add_heading("如果我做的时候卡住了", level=1)
    add_bullets(
        doc,
        [
            "先不要继续搜新工具，先看是不是目标写得太大。",
            "先不要怀疑自己不会用 AI，先看是不是没有把任务拆成步骤。",
            "先不要一口气全改，先只改最卡的一步。",
            "先把一个版本跑完，再去判断下一轮改哪里。",
        ],
    )
    add_paragraph(doc, "很多问题不是出在能力不够，而是出在一开始就把任务做得太复杂。只要我愿意先把问题缩小，很多卡点都会马上松开。")


def add_code_block(doc: Document, title: str, content: str) -> None:
    doc.add_paragraph(title)
    for line in content.strip().splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def add_markdown_examples(doc: Document, folder: str, title: str) -> None:
    doc.add_heading("实际可复制示例", level=1)
    if folder == "06_赠送Skills与Workflow包":
        add_code_block(
            doc,
            "Skill 示例（Markdown）",
            f"""
---
name: {title}
description: 用来处理一类固定任务的简化技能卡
---

## 目标
我先明确这次要完成的结果，而不是直接让 AI 自由发挥。

## 输入
- 服务对象
- 当前素材
- 最终输出格式

## 执行步骤
1. 先拆任务
2. 再列结构
3. 再补表达
4. 最后人工确认
""",
        )
        add_code_block(
            doc,
            "Workflow 示例（Markdown）",
            f"""
# {title} Workflow

1. 明确对象是谁
2. 明确这次只解决一个问题
3. 先写主观点
4. 让 AI 补提纲或表达
5. 改成适合平台的版本
6. 保留一个单一 CTA
""",
        )
    elif folder == "08_可直接复制使用包":
        add_code_block(
            doc,
            "直接复制模板",
            f"""
# {title}

## 服务对象
-

## 当前问题
-

## 我这次想交付的结果
-

## 我准备用的结构
1.
2.
3.

## 最后承接动作
-
""",
        )
    elif folder == "09_模板填空版":
        add_code_block(
            doc,
            "填空模板",
            f"""
# {title}

我服务的人是：_____
他们现在最卡的是：_____
我这次只解决的问题是：_____
我准备给出的结果是：_____
我最后要引导的动作是：_____
""",
        )
    elif folder == "10_训练营作业与打卡版":
        add_code_block(
            doc,
            "作业与打卡模板",
            f"""
# {title}

## 今天必须完成
-

## 今日提交内容
-

## 我今天卡住的点
-

## 我明天只改一个地方
-
""",
        )
    elif folder == "11_成交与预热文案包":
        add_code_block(
            doc,
            "销售页文案骨架",
            f"""
# {title}

你是不是也遇到过这样的情况：
-

为什么很多人一直做不成：
-

我这次准备怎么帮你缩短路径：
-

如果你想先看完整内容，回复我：资料
""",
        )
    else:
        add_code_block(
            doc,
            "实操记录模板",
            f"""
# {title}

## 我这次处理的真实任务
-

## 我保留人工判断的地方
-

## 我交给 AI 提速的地方
-

## 我下一轮准备优化的地方
-
""",
        )


def add_direct_use_section(doc: Document, folder: str, title: str) -> None:
    doc.add_heading("可直接复制使用", level=1)
    if folder == "09_模板填空版":
        blocks = [
            f"模板标题：{title}\n服务对象：_____\n当前痛点：_____\n最想要的结果：_____\n这份内容准备解决的问题：_____",
            "内容骨架：\n1. 先讲当前问题\n2. 再讲为什么很多人做不成\n3. 再讲正确路径\n4. 再给一个最小动作\n5. 最后给一个单一 CTA",
            "提交/承接动作：\n用户看完后要做什么：_____\n你希望他联系你的关键词：_____\n你准备如何跟进：_____",
        ]
    elif folder == "10_训练营作业与打卡版":
        blocks = [
            f"今日作业：{title}\n今天必须完成的 1 个动作：_____\n提交格式：文字 / 图片 / 链接 / 文档\n截止时间：_____",
            "打卡格式：\n我今天完成了：_____\n我卡住的地方：_____\n我明天准备继续推进的动作：_____",
            "复盘问题：\n1. 今天最有效的一步是什么？\n2. 今天最浪费时间的一步是什么？\n3. 明天只改一个点，会改哪里？",
        ]
    elif folder == "11_成交与预热文案包":
        blocks = [
            f"文案骨架：\n开头先讲一个真实问题\n中间讲为什么多数人一直做不成\n再讲你这套内容怎么帮他缩短路径\n最后给一个明确动作",
            "常用 CTA：\n想拿资料，回复我“资料”\n想看课表，回复我“课程”\n想问适不适合自己，回复我“适合吗”",
            "成交回应骨架：\n先回答顾虑\n再强调结果路径\n再提醒时间节点或名额\n最后给一个明确下一步",
        ]
    else:
        blocks = [
            "先把这份内容改成适合你自己业务的话术。",
            "不要整段照搬，先保留结构，再换成你的对象、场景和结果。",
            "每次使用后都记下哪一句最有效，下一轮直接优化这一句。",
        ]
    for block in blocks:
        add_paragraph(doc, block)


def add_folder_specific_deep_dive(doc: Document, folder: str, title: str) -> None:
    if folder == "12_SOP标准操作课":
        doc.add_heading("我会怎么把它写成 SOP", level=1)
        add_numbered(
            doc,
            [
                "先写这件事的触发条件。比如什么时候需要做这一步，而不是心血来潮才做。",
                "再写输入物。比如我要先准备哪些资料、截图、标题草稿、历史案例。",
                "再写执行动作，而且每一步尽量只写一个动作，不要把三个判断塞进一句话里。",
                "最后写输出物。比如我要交付一条标题、一个提纲、一份资料包，而不是写成空泛的‘完成内容优化’。",
            ],
        )
        add_paragraph(doc, "我做 SOP 时最怕的不是步骤少，而是步骤虚。只要别人读完还不知道第一步该动什么，这个 SOP 就还不能交付。")
    elif folder == "13_Workflow设计课":
        doc.add_heading("我会怎么把它串成 Workflow", level=1)
        add_paragraph(doc, "我不会直接画一张很复杂的流程图。我会先把最短路径写出来：入口是什么，中间要经过哪几步，最后的结果是什么。只要这条最短路径跑通，后面再加分支和自动化都不迟。")
        add_bullets(
            doc,
            [
                "入口：这一轮动作是从哪里开始的。",
                "中段：哪一步是 AI 提速，哪一步是我自己拍板。",
                "出口：这一轮结束后留下了什么成品。",
                "反馈：这一轮结果回来以后，我下一轮改哪一步。",
            ],
        )
    elif folder == "14_MCP与工具接入课":
        doc.add_heading("我会怎么理解 MCP", level=1)
        add_paragraph(doc, "我会把 MCP 理解成一根标准插线板。模型本身会思考，但它要想碰到浏览器、文档、表格、网页、数据库这些真实能力，就要通过稳定接口去接。这样我不用每次都从零讲工具怎么调用，而是直接把能力插进流程。")
        add_paragraph(doc, "所以我学 MCP，不是为了研究概念，而是为了让‘会想’和‘会做’之间少一层断裂。只要接得对，我就能把同一套工作方法迁移到不同任务里。")
    elif folder == "15_交付与排错实战课":
        doc.add_heading("我会怎么做最后一公里", level=1)
        add_paragraph(doc, "我会先要求自己交出一个‘现在就能发’的版本。哪怕它还不够满，也必须已经能发、能讲、能承接。因为只有交付出去，我才知道客户真正卡在哪、哪些地方需要补、哪些表达只是我自以为重要。")
        add_paragraph(doc, "排错也是一样。我不会一出问题就全部推翻重来，而是先锁定：问题出在内容、流程、配置还是动作顺序。只要范围锁住，排错速度会快很多。")
    else:
        doc.add_heading("我会怎么把这节课落到真实动作里", level=1)
        add_paragraph(doc, "我会先挑一个最近 3 天一定要做的任务来练，不会挑一个以后也许会做的大项目。因为真正能让我学会的，从来不是理解得多，而是拿一个手头任务反复改顺。")
        add_paragraph(doc, "如果这节课里有方法、有模板、有案例，我会优先保留方法和判断标准，再决定哪些句子、哪些格式适合直接套进我的业务。")


def add_example_assets(doc: Document, folder: str, title: str) -> None:
    doc.add_heading("现成资产示例", level=1)
    if folder == "12_SOP标准操作课":
        add_code_block(
            doc,
            "SOP 模板（Markdown）",
            f"""
# {title} SOP

## 触发条件
- 我要开始处理这类任务时

## 输入物
- 历史案例 3 条
- 当前主题 1 个
- 目标动作 1 个

## 执行步骤
1. 写服务对象
2. 写痛点场景
3. 写主判断
4. 写结构提纲
5. 写交付格式

## 输出物
- 可直接发出的初稿 1 版
""",
        )
    elif folder == "13_Workflow设计课":
        add_code_block(
            doc,
            "Workflow 图文版（Markdown）",
            f"""
# {title} Workflow

入口：收到一个真实任务
1. 先明确这次只解决一个问题
2. 整理输入素材
3. 我自己先定主观点
4. 让 AI 输出提纲或改写版本
5. 我确认最终版本
6. 发布并记录反馈
出口：留下一个可复用版本
""",
        )
    elif folder == "14_MCP与工具接入课":
        add_code_block(
            doc,
            "MCP 配置片段",
            """
{
  "mcpServers": {
    "browser-helper": {
      "command": "npx",
      "args": ["-y", "agent-browser"]
    },
    "doc-helper": {
      "command": "python",
      "args": ["tools/doc_worker.py"]
    }
  }
}
""",
        )
        add_code_block(
            doc,
            "我会怎么接进业务",
            """
1. 浏览器能力负责打开页面和保存草稿
2. 文档能力负责整理内容和导出文档
3. 人工确认负责最后发布和发送
""",
        )
    elif folder == "15_交付与排错实战课":
        add_code_block(
            doc,
            "排错记录模板",
            f"""
# {title} 排错记录

## 问题现象
-

## 我先锁定的范围
- 内容
- 流程
- 配置
- 权限

## 我已经排除的变量
-

## 最后结论
-
""",
        )
    else:
        add_code_block(
            doc,
            "我会保留的实操记录",
            f"""
# {title}

真实任务：
我保留人工判断的地方：
我交给 AI 提速的地方：
我下一轮准备怎么改：
""",
        )


def add_troubleshooting_examples(doc: Document, folder: str) -> None:
    doc.add_heading("出问题时我会怎么查", level=1)
    if folder == "14_MCP与工具接入课":
        add_bullets(
            doc,
            [
                "先看路径对不对。很多接入问题不是能力不行，而是根本没连上正确脚本或配置文件。",
                "再看权限够不够。能读不能写、能打开不能保存，这类问题往往是权限边界没处理好。",
                "再看参数有没有对齐。命令、模型、文件名、输出目录只要有一个写错，结果就会完全跑偏。",
            ],
        )
    elif folder in {"12_SOP标准操作课", "13_Workflow设计课"}:
        add_bullets(
            doc,
            [
                "如果别人照着做还是不会，先看步骤是不是写得太抽象。",
                "如果做的人总在中间卡住，先看是不是某一步输入物没定义清楚。",
                "如果每次结果都不稳定，先看是不是把人工判断和 AI 自动步骤混在了一起。",
            ],
        )
    else:
        add_bullets(
            doc,
            [
                "如果内容一改就散，先回到对象和场景，而不是继续加句子。",
                "如果流程一跑就乱，先删步骤，保留最短路径。",
                "如果客户看了没动作，先查 CTA 是不是写了多个方向。",
            ],
        )


def add_customer_facing_use(doc: Document, folder: str, title: str) -> None:
    doc.add_heading("我会怎么直接发给客户", level=1)
    if folder == "11_成交与预热文案包":
        add_paragraph(doc, "我不会一次把整套文案全丢给客户。我会先发最切他当前问题的那一段，让他先觉得‘这就是我现在卡的地方’，然后再顺势发下一段。这样阅读阻力最小，回应率也更高。")
    elif folder == "14_MCP与工具接入课":
        add_paragraph(doc, "如果客户还处在刚理解阶段，我不会先讲太多配置细节。我会先用一个具体场景去解释，比如‘为什么浏览器能力能帮我把草稿保存到固定位置’，等他觉得场景有用，再补接口和配置。")
    else:
        add_paragraph(doc, f"我直接发这节《{title}》时，会先加一句很短的引导：‘你先看这一份，不用全学，先把这里面一个动作跑顺。’ 这样对方不会被资料量吓住，也更愿意马上动。")


def build_doc(group: dict, doc_item: tuple[str, str, list[str]]) -> Path:
    stem, promise, bullets = doc_item
    path = OUT_DIR / group["folder"] / f"{stem}.docx"
    title = stem.split("_", 1)[1]

    doc = Document()
    configure_document(doc)
    add_title_block(doc, title, group["title"], promise)
    add_paragraph(doc, "我把这一节写成了能直接拿去看、拿去改、拿去用的版本。你不用先把全部学完，先抓住一个真实任务照着跑就行。")

    topic_banner = generate_topic_banner(group, stem, title, promise, bullets)
    if topic_banner.exists():
        add_image(doc, topic_banner, "本篇主题图")
    fallback_key = GROUP_FALLBACK_IMAGE[group["folder"]]
    fallback_path = ROOT / "deliverables" / "course_materials" / "images" / LOCAL_IMAGE_MAP[fallback_key]
    if fallback_path.exists():
        add_image(doc, fallback_path, "辅助理解图")
    local_banner = LOCAL_DIR / f"{group['folder']}.png"
    if local_banner.exists():
        add_image(doc, local_banner, "本模块课程导图")

    build_story_sections(doc, title, promise, bullets)
    add_case_study(doc, title)
    add_real_use_case(doc, group["folder"], title)
    add_folder_specific_deep_dive(doc, group["folder"], title)
    add_how_to_modify(doc, group["folder"], title)
    add_common_questions(doc, title)
    add_plain_examples(doc, title)
    add_problem_solving(doc)
    add_troubleshooting_examples(doc, group["folder"])
    add_practical_steps(doc, title)
    add_example_assets(doc, group["folder"], title)
    add_markdown_examples(doc, group["folder"], title)
    add_direct_use_section(doc, group["folder"], title)
    add_customer_facing_use(doc, group["folder"], title)
    add_takeaway_checklist(doc, title)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("先把一件小事跑顺，再去追求更大的自动化和更复杂的结果。")
    run.font.size = Pt(9)

    doc.save(path)
    return path


def build_group_index(group: dict, docs: list[Path]) -> None:
    doc = Document()
    configure_document(doc)
    add_title_block(doc, group["title"], group["subtitle"], "这个模块不是给你增加更多零散资料，而是帮你形成一条从看懂到会做的连续学习路径。")
    generated = get_generated_cover(group["cover_prefix"])
    if generated:
        add_image(doc, generated, "模块主视觉")
    add_paragraph(doc, "使用建议：不要一次把 5 份全部发给用户。更好的做法是先选 1 到 2 份切中他当前问题的内容发过去，让他先觉得‘这件事我现在就能开始做’，再顺势承接后续内容。")
    doc.add_heading("本模块包含什么", level=1)
    add_bullets(doc, [p.stem for p in docs])
    doc.add_heading("最适合的承接方式", level=1)
    add_bullets(doc, ["可以做免费资料包。", "可以组合成低客单专题。", "可以升级成系统课的一个模块。", "也可以作为训练营开营前的预读材料。"])
    doc.save(OUT_DIR / group["folder"] / "_本模块总目录.docx")


def build_root_index(all_docs: list[tuple[str, Path]]) -> None:
    doc = Document()
    configure_document(doc)
    add_title_block(doc, "AI课程资料库 v3", "以系统课程而不是资料堆砌的方式重组", "结构上分为基础认知、实操课、变现实操课和赠送包，整体更像一套完整产品。")
    add_paragraph(doc, "这一版不再是零散资料集合，而是按客户学习路径重组后的课程材料。前面先帮用户看懂，后面再让用户开始实操，最后再给他工具包和 workflow 包去真正动手。")
    doc.add_heading("总目录", level=1)
    add_bullets(doc, [f"{folder}：{path.stem}" for folder, path in all_docs])
    doc.add_heading("你可以怎么卖和怎么发", level=1)
    add_numbered(doc, ["从基础认知课里挑 1 到 2 份做免费资料。", "把实操课做成低客单小专题。", "把私域转化实操课和训练营逻辑串起来做主产品。", "把赠送包当作成交后的立刻可用资产。"])
    doc.save(OUT_DIR / "00_AI课程资料库v3总目录.docx")


def main() -> None:
    ensure_dirs()
    build_diagram_concept_map()
    build_diagram_content_flow()
    build_diagram_funnel()
    build_diagram_learning_path()
    for group in COURSE_GROUPS:
        generate_local_banner(group)

    all_docs: list[tuple[str, Path]] = []
    for group in COURSE_GROUPS:
        built = []
        for item in group["docs"]:
            p = build_doc(group, item)
            built.append(p)
            all_docs.append((group["folder"], p))
        build_group_index(group, built)
    build_root_index(all_docs)

    with open(OUT_DIR / "README.txt", "w", encoding="utf-8") as fh:
        fh.write("AI 课程资料库 v3 已生成。\n")
        for group in COURSE_GROUPS:
            fh.write(f"\n{group['folder']}\n")
            for item in group["docs"]:
                fh.write(f"- {item[0]}.docx\n")
            fh.write("- _本模块总目录.docx\n")


if __name__ == "__main__":
    main()
